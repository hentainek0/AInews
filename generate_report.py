from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

def add_title(doc, text, level=1):
    if level == 1:
        style = 'Heading 1'
    elif level == 2:
        style = 'Heading 2'
    elif level == 3:
        style = 'Heading 3'
    else:
        style = 'Heading 4'
    
    title = doc.add_heading(text, level=level)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = '微软雅黑'
    title.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return title

def add_paragraph(doc, text, font_size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, data, headers=None):
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if headers and i == 0:
                cell.text = headers[j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                cell.text = data[i - 1 if headers else i][j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    table.style = 'Table Grid'
    return table

def add_code_block(doc, code, language='python'):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    p.paragraph_format.left_indent = Cm(0.74)
    return p

def add_list(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        if ordered:
            p = doc.add_paragraph(f"{i}. {item}", style='List Number')
        else:
            p = doc.add_paragraph(f"• {item}", style='List Bullet')
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def generate_report():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    add_title(doc, 'AI新闻爬虫系统完整分析报告', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(f'报告生成时间：{datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('报告版本：v2.0（包含前端实时动态分析功能）')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    add_title(doc, '目录', level=1)
    toc_items = [
        ('一、系统概述', '1'),
        ('二、项目结构分析', '5'),
        ('三、后端核心模块详解', '10'),
        ('四、前端实时动态分析系统', '35'),
        ('五、Flask后端API设计', '60'),
        ('六、系统流程分析', '75'),
        ('七、技术栈分析', '82'),
        ('八、部署与打包', '90'),
        ('九、系统运行分析', '98'),
        ('十、优缺点评估', '105'),
        ('十一、扩展建议', '112'),
        ('十二、总结', '118'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}'.ljust(40) + f'第 {page} 页')
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    add_title(doc, '一、系统概述', level=1)
    
    add_title(doc, '1.1 系统定位', level=2)
    add_paragraph(doc, '本系统是一个基于 Python 后端与 React 前端的 AI 新闻自动化采集、分析与推送系统。系统从国际主流科技媒体自动抓取 AI 相关新闻，经过摘要提取、中英文翻译处理后，生成结构化的新闻简报，并通过企业微信、钉钉、飞书等即时通讯工具推送到用户移动端。同时，系统提供基于 Web 的前端界面，实现新闻数据的实时动态分析和可视化展示。', indent=True)
    
    data = [
        ['属性', '描述'],
        ['系统名称', 'AI News Crawler and Brief Generator'],
        ['开发语言', 'Python 3.13 + TypeScript'],
        ['前端框架', 'React 18 + Vite 6'],
        ['后端框架', 'Flask 3'],
        ['运行平台', 'Windows 11'],
        ['核心功能', '新闻采集、文本摘要、机器翻译、消息推送、实时动态分析'],
        ['目标用户', 'AI 从业者、科技爱好者、企业团队、数据分析人员'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '1.2 系统架构层次', level=2)
    add_paragraph(doc, '本系统采用前后端分离的分层架构设计，各层职责清晰，便于维护和扩展：', indent=True)
    
    architecture = """┌───────────────────────────────────────────────────────────────────────┐
│                        前端展示层 (Frontend Layer)                     │
│  React 18 + TypeScript + TailwindCSS + Chart.js                       │
│  Dashboard | NewsList | Analytics | Settings                          │
├───────────────────────────────────────────────────────────────────────┤
│                        后端API层 (Backend API Layer)                   │
│  Flask 3 + Flask-CORS                                                 │
│  /api/news | /api/analytics | /api/config | /api/config/schedule      │
├───────────────────────────────────────────────────────────────────────┤
│                        业务逻辑层 (Business Logic Layer)               │
│  crawl_ai_news() | generate_brief() | translate() | PushNotifier      │
│  extract_summary() | extract_tags()                                    │
├───────────────────────────────────────────────────────────────────────┤
│                        数据服务层 (Data Service Layer)                  │
│  requests | BeautifulSoup | translators | json | datetime             │
├───────────────────────────────────────────────────────────────────────┤
│                        数据源 (Data Sources)                           │
│  TechCrunch AI RSS | VentureBeat AI RSS                               │
│  企业微信/钉钉/飞书 API | articles.json | config.json | schedule.json  │
└───────────────────────────────────────────────────────────────────────┘"""
    add_code_block(doc, architecture, 'text')
    
    add_title(doc, '1.3 系统特性', level=2)
    features = [
        '多源新闻采集：支持 TechCrunch AI、VentureBeat AI 等多个 RSS 数据源',
        '智能摘要提取：基于词频统计的文本摘要算法，自动提取关键信息',
        '多层翻译降级：支持百度、Google、腾讯翻译及 MyMemory API 多层降级',
        '多平台消息推送：支持企业微信、钉钉、飞书即时通讯工具推送',
        '定时任务调度：基于标准库实现定时任务，支持每天固定时间运行',
        '实时动态分析：前端可视化展示新闻趋势、来源分布、关键词统计',
        '数据持久化：JSON 文件存储新闻数据和配置信息',
        '一键打包部署：PyInstaller 打包为单文件，零依赖运行',
        '前后端分离：RESTful API 设计，支持多端调用',
        '响应式界面：深蓝色主题，支持移动端和桌面端访问',
    ]
    add_list(doc, features)
    
    doc.add_page_break()
    
    add_title(doc, '二、项目结构分析', level=1)
    
    add_title(doc, '2.1 完整目录结构', level=2)
    dir_structure = """d:\\ruangong233\\demo2\\news\\
├── __pycache__/                          # Python编译缓存
├── api/                                  # Flask后端API
│   ├── __pycache__/
│   └── app.py                            # Flask应用入口
├── data/                                 # 数据存储目录
│   ├── articles.json                     # 新闻文章数据
│   ├── config.json                       # 系统配置
│   └── schedule.json                     # 定时任务配置
├── frontend/                             # React前端项目
│   ├── public/
│   │   └── favicon.svg
│   ├── src/
│   │   ├── api/
│   │   │   └── index.ts                  # API调用封装
│   │   ├── components/
│   │   │   ├── Empty.tsx                 # 空状态组件
│   │   │   └── Layout.tsx                # 布局组件
│   │   ├── pages/
│   │   │   ├── Dashboard.tsx             # 数据概览页面
│   │   │   ├── NewsList.tsx              # 新闻列表页面
│   │   │   ├── Analytics.tsx             # 数据分析页面
│   │   │   └── Settings.tsx              # 系统设置页面
│   │   ├── types/
│   │   │   └── index.ts                  # TypeScript类型定义
│   │   ├── App.tsx                       # 主应用组件
│   │   ├── index.css                     # 全局样式
│   │   └── main.tsx                      # 入口文件
│   ├── .gitignore
│   ├── README.md
│   ├── eslint.config.js
│   ├── index.html
│   ├── package-lock.json
│   ├── package.json
│   ├── postcss.config.js
│   ├── tailwind.config.js
│   ├── tsconfig.json
│   └── vite.config.ts
├── AI_News_Brief_*.md                    # 生成的新闻简报文件
├── AI_新闻爬虫系统分析报告.docx           # 系统分析报告
├── README.md                             # 项目说明文档
├── debug_main.py                         # 调试入口
├── generate_report.py                    # Word报告生成脚本
├── main.py                               # 完整版本入口
├── md_to_pdf.py                          # Markdown转PDF工具
├── news_crawler.py                       # 新闻爬虫模块
├── news_scheduler.py                     # 定时任务调度器
├── news_summary.py                       # 新闻摘要模块
├── push_notifier.py                      # 推送通知模块
├── requirements.txt                      # Python依赖清单
├── run_once.py                           # 单次运行版本
└── simple_main.py                        # 简化版核心入口"""
    add_code_block(doc, dir_structure, 'text')
    
    add_title(doc, '2.2 文件职责说明', level=2)
    data = [
        ['文件', '职责', '所属层次'],
        ['simple_main.py', '核心入口，包含爬取、摘要、翻译、简报生成', '业务逻辑层'],
        ['news_crawler.py', '新闻爬虫类，封装RSS采集逻辑', '业务逻辑层'],
        ['news_summary.py', '文本摘要类，基于NLTK的摘要提取', '业务逻辑层'],
        ['push_notifier.py', '推送通知类，支持多平台消息推送', '业务逻辑层'],
        ['news_scheduler.py', '定时任务调度器，标准库实现', '业务逻辑层'],
        ['api/app.py', 'Flask后端API，提供RESTful接口', '后端API层'],
        ['frontend/src/App.tsx', 'React主应用，路由管理', '前端展示层'],
        ['frontend/src/pages/Dashboard.tsx', '数据概览页面', '前端展示层'],
        ['frontend/src/pages/NewsList.tsx', '新闻列表页面', '前端展示层'],
        ['frontend/src/pages/Analytics.tsx', '数据分析页面', '前端展示层'],
        ['frontend/src/pages/Settings.tsx', '系统设置页面', '前端展示层'],
        ['frontend/src/api/index.ts', 'API调用封装', '前端展示层'],
        ['data/articles.json', '新闻文章数据存储', '数据存储'],
        ['data/config.json', '系统配置存储', '数据存储'],
        ['data/schedule.json', '定时任务配置存储', '数据存储'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '三、后端核心模块详解', level=1)
    
    add_title(doc, '3.1 新闻爬取模块 (crawl_ai_news)', level=2)
    add_paragraph(doc, '位置：simple_main.py (第29-73行)', indent=True)
    add_paragraph(doc, '功能：从 RSS 订阅源抓取 AI 新闻文章，支持多源采集、防反爬机制和容错处理', indent=True)
    
    code = """def crawl_ai_news():
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    rss_feeds = [
        {'name': 'TechCrunch AI', 'url': 'https://techcrunch.com/tag/ai/feed/'},
        {'name': 'VentureBeat AI', 'url': 'https://venturebeat.com/category/ai/feed/'},
    ]
    
    articles = []
    for feed in rss_feeds:
        print(f"Crawling {feed['name']}...")
        try:
            time.sleep(random.uniform(0.5, 1.5))  # 防反爬延迟
            response = requests.get(feed['url'], headers=headers, timeout=10)
            soup = BeautifulSoup(response.text, 'xml')
            items = soup.find_all('item')[:6]
            for item in items:
                title = item.find('title').get_text(strip=True) if item.find('title') else ''
                link = item.find('link').get_text(strip=True) if item.find('link') else ''
                desc = item.find('description').get_text(strip=True)[:200] if item.find('description') else ''
                if title and link:
                    articles.append({
                        'title': title, 'link': link, 'content': desc + ' ...',
                        'source': feed['name'], 'date': ''
                    })
        except Exception as e:
            print(f"Failed to crawl {feed['name']}: {e}")
    
    # 爬取失败时使用样例数据
    if not articles:
        articles = [{'title': 'AI Breakthrough...', 'link': 'https://example.com/ai1', ...}]
    
    return articles"""
    add_code_block(doc, code)
    
    add_title(doc, '设计特点', level=3)
    data = [
        ['特性', '说明'],
        ['多源采集', '支持 TechCrunch AI 和 VentureBeat AI 两个权威数据源'],
        ['防反爬机制', '随机延迟 0.5-1.5 秒，模拟人类访问行为'],
        ['容错处理', '爬取失败时自动使用预设样例数据，保证系统可用性'],
        ['数据限制', '每个源最多采集 6 篇文章，避免数据过载'],
        ['User-Agent 伪装', '设置浏览器 User-Agent，避免被识别为爬虫'],
        ['超时控制', '设置 10 秒超时，避免长时间阻塞'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '数据结构', level=3)
    code = """article = {
    'title': str,      # 文章标题（英文）
    'link': str,       # 原文链接
    'content': str,    # 文章内容/摘要（英文）
    'source': str,     # 来源名称（如 TechCrunch AI）
    'date': str        # 发布日期（可选）
}"""
    add_code_block(doc, code)
    
    add_title(doc, '3.2 文本摘要模块 (extract_summary)', level=2)
    add_paragraph(doc, '位置：simple_main.py (第75-114行)', indent=True)
    add_paragraph(doc, '功能：基于词频统计的无监督文本摘要提取算法，无需外部依赖，纯 Python 实现', indent=True)
    
    add_title(doc, '算法流程', level=3)
    steps = [
        '句子分割：使用正则表达式按句号、问号分割句子，智能避免误分割缩略词',
        '停用词过滤：移除 57 个常见英文停用词（a, the, is, are 等）',
        '词频计算：统计有效词频并归一化（除以最大词频）',
        '句子评分：根据句子包含的关键词频率计算得分，考虑句子长度加权',
        '排序选取：选取得分最高的前 N 个句子作为摘要',
        '顺序还原：按原文顺序重新排列选中的句子，保持语义连贯',
    ]
    add_list(doc, steps, ordered=True)
    
    add_title(doc, '关键技术点', level=3)
    code = """STOP_WORDS = {'a', 'about', 'above', 'after', 'again', 'against', 'all', 'am', 
              'an', 'and', 'any', 'are', 'as', 'at', 'be', 'because', 'been', 
              'before', 'being', 'below', 'between', 'both', 'but', 'by', ...}
# 共 57 个英文停用词，手动定义避免依赖 NLTK 下载

def extract_summary(text, num_sentences=2):
    # 智能正则：避免误分割如 "Mr. Smith" 这样的缩略词
    sentences = re.split(r'(?<!\\w\\.\\w.)(?<![A-Z][a-z]\\.)(?<=\\.|\\?)\\s', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    word_freq = {}
    for sentence in sentences:
        words = re.findall(r'\\b[a-zA-Z]+\\b', sentence.lower())
        for word in words:
            if word not in STOP_WORDS:
                word_freq[word] = word_freq.get(word, 0) + 1
    
    # 归一化词频
    if word_freq:
        max_freq = max(word_freq.values())
        for word in word_freq:
            word_freq[word] /= max_freq
    
    # 句子评分
    sentence_scores = {}
    for sentence in sentences:
        score = 0
        words = re.findall(r'\\b[a-zA-Z]+\\b', sentence.lower())
        for word in words:
            if word in word_freq:
                score += word_freq[word]
        if len(words) > 0:
            score /= len(words)  # 长度加权
        sentence_scores[sentence] = score"""
    add_code_block(doc, code)
    
    add_title(doc, '算法优缺点分析', level=3)
    data = [
        ['优点', '缺点'],
        ['无外部依赖，纯 Python 标准库实现', '仅支持英文文本，不支持中文摘要'],
        ['基于统计方法，计算效率高（O(n)）', '缺乏语义理解能力，对歧义句处理效果一般'],
        ['可配置摘要长度，灵活适应不同场景', '对短文本（<50字符）处理效果不佳'],
        ['算法简单易懂，易于维护和调试', '不考虑句子之间的语义关联'],
        ['句子顺序还原，保持原文语义连贯', '依赖词频，对低频关键词可能遗漏'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.3 翻译模块 (translate_text)', level=2)
    add_paragraph(doc, '位置：simple_main.py (第116-146行)', indent=True)
    add_paragraph(doc, '功能：将英文文本翻译为中文，采用多层降级策略保证翻译成功率', indent=True)
    
    add_title(doc, '多层降级策略', level=3)
    code = """def translate_text(text, target_lang='zh'):
    if not text or len(text.strip()) == 0:
        return text
    
    try:
        # 第一层：translators库 - 优先使用百度翻译
        import translators as ts
        try:
            return ts.baidu(text, to_language=target_lang)
        except:
            try:
                return ts.google(text, to_language=target_lang)
            except:
                try:
                    return ts.tencent(text, to_language=target_lang)
                except:
                    return text
    except ImportError:
        # 第二层：MyMemory API - 免费翻译服务（无需安装translators库）
        url = "https://api.mymemory.translated.net/get"
        params = {'q': text[:500], 'langpair': 'en|zh'}
        response = requests.get(url, params=params, timeout=10)
        result = response.json()
        if result.get('responseData'):
            return result['responseData']['translatedText']
        return text"""
    add_code_block(doc, code)
    
    add_title(doc, '翻译服务对比', level=3)
    data = [
        ['服务', '特点', '稳定性', '依赖要求'],
        ['百度翻译', '中文支持好，翻译准确率高', '高', '需要安装 translators 库'],
        ['Google翻译', '翻译质量最高，支持多语言', '低（需翻墙）', '需要安装 translators 库'],
        ['腾讯翻译', '多语言支持，响应速度快', '中', '需要安装 translators 库'],
        ['MyMemory', '免费额度有限（1000字符/天）', '中', '无需额外依赖，仅需 requests'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.4 简报生成模块 (generate_brief)', level=2)
    add_paragraph(doc, '位置：simple_main.py (第148-204行)', indent=True)
    add_paragraph(doc, '功能：整合所有文章生成结构化新闻简报，支持中英文双语输出', indent=True)
    
    add_title(doc, '输出格式示例', level=3)
    example = """# AI 新闻简报

## 概览
从多个科技媒体来源收集了 5 篇 AI 相关新闻文章。

## 新闻详情
### 1. Anthropic的新Claude功能正在悄悄地向您推销人工智能
**来源**: TechCrunch AI
**摘要**: Claude的新Reflect仪表板不仅能直观地显示您如何使用人工智能...
**链接**: [https://techcrunch.com/...](https://techcrunch.com/...)

## 重点关注领域
- 最新 AI 技术发展趋势
- AI 在各行业的应用
- AI 伦理与监管话题
- 主要科技公司的 AI 战略

---
Generated: 2026-07-13 10:30:00
来源: TechCrunch, VentureBeat, ZDNet, MIT Technology Review"""
    add_code_block(doc, example, 'markdown')
    
    add_title(doc, '3.5 推送通知模块 (PushNotifier)', level=2)
    add_paragraph(doc, '位置：push_notifier.py (第4-143行)', indent=True)
    add_paragraph(doc, '功能：支持企业微信、钉钉、飞书多平台消息推送，统一接口封装', indent=True)
    
    add_title(doc, '类结构', level=3)
    code = """class PushNotifier:
    def __init__(self, config=None):
        self.config = config or {}
    
    def send_wechat(self, content, webhook_url=None):
        # 企业微信推送 - Markdown格式，支持丰富的文本样式
        # 消息限制：4096字符
    
    def send_dingtalk(self, content, access_token=None):
        # 钉钉推送 - Markdown格式，支持标题、列表等格式
        # 消息限制：4096字符
    
    def send_feishu(self, content, webhook_url=None):
        # 飞书推送 - 纯文本格式（飞书机器人不支持Markdown）
        # 消息限制：3000字符
    
    def format_brief_for_push(self, brief):
        # 将Markdown格式简报转换为适合推送的格式
    
    def send_all(self, brief):
        # 批量推送所有配置的平台，统计成功/失败数量"""
    add_code_block(doc, code)
    
    add_title(doc, '各平台特性对比', level=3)
    data = [
        ['平台', '消息格式', '内容限制', 'URL格式', 'Webhook获取方式'],
        ['企业微信', 'Markdown', '4096字符', 'https://qyapi.weixin.qq.com/cgi-bin/webhook/send?key=xxx', '群聊设置→添加机器人'],
        ['钉钉', 'Markdown', '4096字符', 'https://oapi.dingtalk.com/robot/send?access_token=xxx', '群聊设置→添加自定义机器人'],
        ['飞书', '纯文本', '3000字符', 'https://open.feishu.cn/open-apis/bot/v2/hook/xxx', '群聊设置→添加自定义机器人'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.6 定时任务模块 (news_scheduler)', level=2)
    add_paragraph(doc, '位置：news_scheduler.py (第56-69行)', indent=True)
    add_paragraph(doc, '功能：实现定时任务调度，支持每天固定时间运行，使用 Python 标准库实现，无第三方依赖', indent=True)
    
    add_title(doc, '核心算法', level=3)
    code = """def get_next_run_time(target_times):
    now = datetime.datetime.now()
    today = now.date()
    
    for target_time_str in target_times:
        # 解析目标时间（如 "09:00"）
        target_time = datetime.datetime.strptime(target_time_str, "%H:%M").time()
        target_datetime = datetime.datetime.combine(today, target_time)
        
        # 如果今天的目标时间还未到达，返回该时间
        if target_datetime > now:
            return target_datetime
    
    # 如果今天已过所有时间点，返回明天第一个时间
    first_time = datetime.datetime.strptime(target_times[0], "%H:%M").time()
    next_day = today + datetime.timedelta(days=1)
    return datetime.datetime.combine(next_day, first_time)"""
    add_code_block(doc, code)
    
    add_title(doc, '调度策略', level=3)
    data = [
        ['配置项', '默认值', '说明'],
        ['schedule_times', '["09:00", "18:00"]', '每天运行时间点，支持多个时间'],
        ['wait_interval', '1秒', '轮询间隔，每秒检查一次是否到达运行时间'],
        ['KeyboardInterrupt', '支持', '用户可通过 Ctrl+C 优雅停止'],
        ['异常处理', '自动重试', '调度异常时等待60秒后重新尝试'],
        ['立即执行', '支持', '启动时立即执行一次，然后进入定时循环'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '四、前端实时动态分析系统', level=1)
    
    add_title(doc, '4.1 前端技术栈', level=2)
    data = [
        ['技术', '版本', '用途'],
        ['React', '18', '前端框架，组件化开发'],
        ['TypeScript', '-', '类型安全，提升代码质量'],
        ['Vite', '6', '构建工具，快速热更新'],
        ['TailwindCSS', '3', 'CSS框架，快速样式开发'],
        ['Chart.js', '-', '图表库，数据可视化'],
        ['Lucide React', '-', '图标库，UI图标'],
        ['Axios', '-', 'HTTP客户端，API调用'],
        ['Zustand', '-', '状态管理（可选）'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.2 前端项目结构', level=2)
    structure = """frontend/src/
├── api/
│   └── index.ts                  # API调用封装（newsApi, analyticsApi, configApi）
├── components/
│   ├── Empty.tsx                 # 空状态组件（无数据时展示）
│   └── Layout.tsx                # 布局组件（侧边栏导航 + 主内容区）
├── pages/
│   ├── Dashboard.tsx             # 数据概览页面（统计卡片 + 趋势图表）
│   ├── NewsList.tsx              # 新闻列表页面（搜索 + 筛选 + 分页 + 详情弹窗）
│   ├── Analytics.tsx             # 数据分析页面（趋势图 + 饼图 + 词云）
│   └── Settings.tsx              # 系统设置页面（数据源 + 推送配置 + 定时任务）
├── types/
│   └── index.ts                  # TypeScript类型定义（News, Stats, Config等）
├── App.tsx                       # 主应用组件（路由管理）
├── index.css                     # 全局样式（深色主题 + 滚动条样式）
└── main.tsx                      # 入口文件（ReactDOM渲染）"""
    add_code_block(doc, structure, 'text')
    
    add_title(doc, '4.3 Layout组件设计', level=2)
    add_paragraph(doc, '位置：frontend/src/components/Layout.tsx', indent=True)
    add_paragraph(doc, '功能：统一布局组件，包含侧边栏导航和主内容区域，支持页面切换', indent=True)
    
    code = """function Layout({ children, currentPage, onPageChange }: LayoutProps) {
    const navItems = [
        { id: '/', label: '数据概览', icon: LayoutDashboard },
        { id: '/news', label: '新闻列表', icon: Newspaper },
        { id: '/analytics', label: '数据分析', icon: BarChart3 },
        { id: '/settings', label: '系统设置', icon: Settings },
    ];
    
    return (
        <div className="flex min-h-screen bg-slate-900">
            {/* 侧边栏 */}
            <aside className="w-64 bg-slate-800 border-r border-slate-700">
                <div className="p-6 border-b border-slate-700">
                    <h1 className="text-xl font-bold text-blue-400">AI 新闻爬虫</h1>
                </div>
                <nav className="p-4 space-y-2">
                    {navItems.map(item => (
                        <button
                            key={item.id}
                            onClick={() => onPageChange(item.id)}
                            className={`w-full flex items-center gap-3 px-4 py-3 rounded-lg transition-colors ${
                                currentPage === item.id
                                    ? 'bg-blue-600 text-white'
                                    : 'text-slate-300 hover:bg-slate-700'
                            }`}
                        >
                            <item.icon size={20} />
                            {item.label}
                        </button>
                    ))}
                </nav>
            </aside>
            
            {/* 主内容区 */}
            <main className="flex-1 p-8 overflow-auto">
                {children}
            </main>
        </div>
    );
}"""
    add_code_block(doc, code)
    
    add_title(doc, '4.4 Dashboard页面设计', level=2)
    add_paragraph(doc, '位置：frontend/src/pages/Dashboard.tsx', indent=True)
    add_paragraph(doc, '功能：数据概览页面，展示关键指标统计卡片、近7天新闻趋势图表、来源分布和最新新闻预览', indent=True)
    
    add_title(doc, '统计卡片设计', level=3)
    data = [
        ['指标', '数据来源', '展示方式'],
        ['新闻总数', '/api/analytics/stats → total_count', '数字卡片'],
        ['今日新增', '/api/analytics/stats → today_count', '数字卡片（带增减箭头）'],
        ['来源数量', '/api/analytics/stats → source_distribution.length', '数字卡片'],
        ['热门关键词', '/api/analytics/stats → top_keywords[0]', '标签展示'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '趋势图表设计', level=3)
    code = """// 使用 Chart.js + react-chartjs-2 实现折线图
const trendChartData = {
    labels: weekTrend.map(d => d.date),
    datasets: [{
        label: '新闻数量',
        data: weekTrend.map(d => d.count),
        borderColor: '#3b82f6',
        backgroundColor: 'rgba(59, 130, 246, 0.1)',
        fill: true,
        tension: 0.4,
    }]
};

const chartOptions = {
    responsive: true,
    plugins: {
        legend: { display: false },
        tooltip: {
            backgroundColor: '#1e293b',
            titleColor: '#fff',
            bodyColor: '#94a3b8',
        }
    },
    scales: {
        x: { grid: { color: '#334155' }, ticks: { color: '#94a3b8' } },
        y: { grid: { color: '#334155' }, ticks: { color: '#94a3b8' } },
    }
};"""
    add_code_block(doc, code)
    
    add_title(doc, '4.5 NewsList页面设计', level=2)
    add_paragraph(doc, '位置：frontend/src/pages/NewsList.tsx', indent=True)
    add_paragraph(doc, '功能：新闻列表页面，支持关键词搜索、来源筛选、分页浏览和新闻详情弹窗', indent=True)
    
    add_title(doc, '搜索筛选功能', level=3)
    code = """// 关键词搜索和来源筛选
const [keyword, setKeyword] = useState('');
const [source, setSource] = useState('');

// 调用API获取筛选后的新闻列表
const { data, isLoading } = useQuery(
    ['news', page, keyword, source],
    () => newsApi.getNews(page, limit, source || undefined, keyword || undefined)
);

// 来源选项动态生成
const sourceOptions = [...new Set(articles.map(a => a.source))];"""
    add_code_block(doc, code)
    
    add_title(doc, '分页设计', level=3)
    code = """const totalPages = Math.ceil(total / limit);
const pages = Array.from({ length: totalPages }, (_, i) => i + 1);

// 分页按钮组
<div className="flex items-center justify-center gap-2 mt-8">
    <button onClick={() => setPage(p => Math.max(1, p - 1))} disabled={page === 1}>
        上一页
    </button>
    {pages.map(p => (
        <button
            key={p}
            onClick={() => setPage(p)}
            className={page === p ? 'bg-blue-600' : 'bg-slate-700'}
        >
            {p}
        </button>
    ))}
    <button onClick={() => setPage(p => Math.min(totalPages, p + 1))} disabled={page === totalPages}>
        下一页
    </button>
</div>"""
    add_code_block(doc, code)
    
    add_title(doc, '4.6 Analytics页面设计', level=2)
    add_paragraph(doc, '位置：frontend/src/pages/Analytics.tsx', indent=True)
    add_paragraph(doc, '功能：数据分析页面，包含趋势分析折线图、来源分布饼图和热门关键词词云', indent=True)
    
    add_title(doc, '来源分布饼图', level=3)
    code = """// 使用 Chart.js 饼图展示来源分布
const pieChartData = {
    labels: sourceDistribution.map(d => d.name),
    datasets: [{
        data: sourceDistribution.map(d => d.count),
        backgroundColor: [
            '#3b82f6', // blue
            '#8b5cf6', // purple
            '#ec4899', // pink
            '#f59e0b', // amber
            '#10b981', // emerald
        ],
        borderWidth: 0,
    }]
};

const pieOptions = {
    responsive: true,
    plugins: {
        legend: {
            position: 'bottom' as const,
            labels: { color: '#94a3b8', padding: 20 },
        }
    }
};"""
    add_code_block(doc, code)
    
    add_title(doc, '热门关键词词云', level=3)
    code = """// 关键词云展示，根据词频动态调整字体大小
<div className="flex flex-wrap gap-2">
    {topKeywords.map((kw, index) => {
        const maxCount = Math.max(...topKeywords.map(k => k.count));
        const size = 14 + (kw.count / maxCount) * 16;
        const colors = ['#3b82f6', '#8b5cf6', '#ec4899', '#f59e0b', '#10b981'];
        return (
            <span
                key={index}
                style={{ fontSize: `${size}px`, color: colors[index % colors.length] }}
                className="bg-slate-800 px-3 py-1 rounded-full"
            >
                {kw.word} ({kw.count})
            </span>
        );
    })}
</div>"""
    add_code_block(doc, code)
    
    add_title(doc, '4.7 Settings页面设计', level=2)
    add_paragraph(doc, '位置：frontend/src/pages/Settings.tsx', indent=True)
    add_paragraph(doc, '功能：系统设置页面，包含 RSS 数据源配置、推送配置（企业微信/钉钉/飞书）、定时任务配置和推送测试功能', indent=True)
    
    add_title(doc, '配置管理功能', level=3)
    data = [
        ['配置项', '功能', '持久化方式'],
        ['RSS 数据源', '添加/删除/编辑数据源，支持自定义源', 'config.json'],
        ['企业微信 Webhook', '配置企业微信机器人地址', 'config.json'],
        ['钉钉 Access Token', '配置钉钉机器人令牌', 'config.json'],
        ['飞书 Webhook', '配置飞书机器人地址', 'config.json'],
        ['定时任务时间', '配置每天运行时间点，支持多个时间', 'schedule.json'],
        ['任务状态', '查看上次运行时间、下次运行时间、启用状态', 'schedule.json'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '推送测试功能', level=3)
    code = """const handleTestPush = async () => {
    setTestResult('测试中...');
    try {
        await configApi.testPush({
            wechat_webhook: config.wechat_webhook,
            dingtalk_token: config.dingtalk_token,
            feishu_webhook: config.feishu_webhook,
        });
        setTestResult('推送测试成功');
    } catch (error) {
        setTestResult('推送测试失败');
    }
    setTimeout(() => setTestResult(''), 3000);
};"""
    add_code_block(doc, code)
    
    add_title(doc, '4.8 API调用封装', level=2)
    add_paragraph(doc, '位置：frontend/src/api/index.ts', indent=True)
    add_paragraph(doc, '功能：统一封装 API 调用，使用 Axios 创建实例，设置基础 URL 和超时时间', indent=True)
    
    code = """import axios from 'axios';

const API_BASE = 'http://localhost:5000/api';

const api = axios.create({
    baseURL: API_BASE,
    timeout: 30000,
});

export const newsApi = {
    getNews: (page=1, limit=10, source?, keyword?) =>
        api.get('/news', { params: { page, limit, source, keyword } }),
    crawlNews: () => api.post('/news'),
    deleteNews: (id) => api.delete(`/news/${id}`),
};

export const analyticsApi = {
    getStats: () => api.get('/analytics/stats'),
};

export const configApi = {
    getConfig: () => api.get('/config'),
    updateConfig: (data) => api.put('/config', data),
    testPush: (data) => api.post('/config/push', data),
    getSchedule: () => api.get('/config/schedule'),
    updateSchedule: (data) => api.put('/config/schedule', data),
};"""
    add_code_block(doc, code)
    
    doc.add_page_break()
    
    add_title(doc, '五、Flask后端API设计', level=1)
    
    add_title(doc, '5.1 API架构设计', level=2)
    add_paragraph(doc, '位置：api/app.py', indent=True)
    add_paragraph(doc, '功能：基于 Flask 框架构建的 RESTful API，提供新闻管理、数据分析和配置管理接口', indent=True)
    
    add_title(doc, 'API接口列表', level=3)
    data = [
        ['接口', '方法', '功能', '参数', '返回数据'],
        ['/api/news', 'GET', '获取新闻列表', 'page, limit, source, keyword', '新闻数组 + 分页信息'],
        ['/api/news', 'POST', '爬取新闻', '无', '爬取结果 + 文章数量'],
        ['/api/news/:id', 'DELETE', '删除新闻', 'article_id', '操作结果'],
        ['/api/analytics/stats', 'GET', '获取统计数据', '无', '总数、趋势、来源分布、关键词'],
        ['/api/config', 'GET', '获取系统配置', '无', '配置对象'],
        ['/api/config', 'PUT', '更新系统配置', '配置对象', '操作结果'],
        ['/api/config/push', 'POST', '测试推送', 'webhook配置', '测试结果'],
        ['/api/config/schedule', 'GET', '获取定时任务配置', '无', '定时任务对象'],
        ['/api/config/schedule', 'PUT', '更新定时任务配置', '定时任务对象', '操作结果'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.2 数据存储设计', level=2)
    data = [
        ['文件', '结构', '用途'],
        ['articles.json', '[{id, title, title_zh, link, summary, summary_zh, content, source, date, tags, created_at}]', '存储新闻文章数据，最多100条'],
        ['config.json', '{wechat_webhook, dingtalk_token, feishu_webhook, rss_sources, article_count, auto_translate}', '存储系统配置'],
        ['schedule.json', '{times, enabled, last_run, next_run}', '存储定时任务配置和运行状态'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.3 数据加载与保存', level=2)
    code = """def load_articles():
    if os.path.exists(ARTICLES_FILE):
        with open(ARTICLES_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_articles(articles):
    with open(ARTICLES_FILE, 'w', encoding='utf-8') as f:
        json.dump(articles, f, indent=2, ensure_ascii=False)

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    # 返回默认配置
    return {
        'wechat_webhook': '',
        'dingtalk_token': '',
        'feishu_webhook': '',
        'rss_sources': [
            {'name': 'TechCrunch AI', 'url': 'https://techcrunch.com/tag/ai/feed/'},
            {'name': 'VentureBeat AI', 'url': 'https://venturebeat.com/category/ai/feed/'}
        ],
        'article_count': 5,
        'auto_translate': True
    }"""
    add_code_block(doc, code)
    
    add_title(doc, '5.4 新闻爬取API实现', level=2)
    code = """@app.route('/api/news', methods=['POST'])
def crawl_and_save_news():
    try:
        config = load_config()
        articles = crawl_ai_news()
        
        # 处理文章数据，添加ID、翻译字段、标签等
        translated_articles = []
        for i, article in enumerate(articles):
            article_id = f"{datetime.now().strftime('%Y%m%d')}_{str(i+1).zfill(3)}"
            translated_articles.append({
                'id': article_id,
                'title': article['title'],
                'title_zh': article.get('title_zh', article['title']),
                'link': article['link'],
                'summary': article['content'][:200],
                'summary_zh': article.get('summary_zh', article['content'][:200]),
                'content': article['content'],
                'source': article['source'],
                'date': article.get('date', datetime.now().strftime('%Y-%m-%d')),
                'tags': extract_tags(article['title']),
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
        
        # 合并到现有数据，保留最新100条
        existing = load_articles()
        all_articles = translated_articles + existing
        save_articles(all_articles[:100])
        
        # 生成简报并推送
        brief = generate_brief(articles[:config.get('article_count', 5)])
        if config.get('wechat_webhook') or config.get('dingtalk_token') or config.get('feishu_webhook'):
            notifier = PushNotifier(config)
            notifier.send_all(brief)
        
        # 更新调度状态
        schedule = load_schedule()
        schedule['last_run'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        save_schedule(schedule)
        
        return jsonify({'status': 'success', 'message': f'Successfully crawled {len(articles)} articles', 'count': len(articles)})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500"""
    add_code_block(doc, code)
    
    add_title(doc, '5.5 统计数据API实现', level=2)
    code = """@app.route('/api/analytics/stats', methods=['GET'])
def get_stats():
    articles = load_articles()
    
    # 今日新增数量
    today = datetime.now().strftime('%Y-%m-%d')
    today_count = sum(1 for a in articles if a['date'] == today)
    
    # 来源分布
    source_distribution = {}
    for a in articles:
        source_distribution[a['source']] = source_distribution.get(a['source'], 0) + 1
    
    # 关键词统计
    keywords = {}
    for a in articles:
        for tag in a.get('tags', []):
            keywords[tag] = keywords.get(tag, 0) + 1
        # 从标题提取关键词
        words = re.findall(r'\\b[A-Za-z]{3,}\\b', a['title'].lower())
        for word in words:
            if word not in STOP_WORDS:
                keywords[word] = keywords.get(word, 0) + 1
    
    top_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)[:10]
    
    # 近7天趋势
    week_trend = []
    for i in range(7):
        date = (datetime.now() - timedelta(days=i)).strftime('%Y-%m-%d')
        count = sum(1 for a in articles if a['date'] == date)
        week_trend.append({'date': date, 'count': count})
    week_trend = week_trend[::-1]
    
    return jsonify({
        'status': 'success',
        'data': {
            'total_count': len(articles),
            'today_count': today_count,
            'week_trend': week_trend,
            'source_distribution': [{'name': k, 'count': v} for k, v in source_distribution.items()],
            'top_keywords': [{'word': k, 'count': v} for k, v in top_keywords]
        }
    })"""
    add_code_block(doc, code)
    
    add_title(doc, '5.6 标签提取功能', level=2)
    code = """def extract_tags(title):
    # 预定义的AI相关关键词列表
    keywords = ['AI', 'Artificial Intelligence', 'Machine Learning', 'Deep Learning', 
                'Neural Network', 'GPT', 'LLM', 'Chatbot', 'Automation', 'Robotics', 
                'Computer Vision', 'Natural Language', 'NLP', 'Data Science', 
                'Big Data', 'Cloud', 'Algorithm']
    
    tags = []
    for kw in keywords:
        if kw.lower() in title.lower():
            tags.append(kw)
    return tags[:3]  # 最多提取3个标签"""
    add_code_block(doc, code)
    
    doc.add_page_break()
    
    add_title(doc, '六、系统流程分析', level=1)
    
    add_title(doc, '6.1 单次运行流程', level=2)
    flow = """┌─────────────┐    ┌─────────────┐    ┌─────────────┐    ┌─────────────┐    ┌─────────────┐
│  Step 1     │ -> │  Step 2     │ -> │  Step 3     │ -> │  Step 4     │ -> │   Complete  │
│  爬取新闻   │    │  生成简报   │    │  保存文件   │    │  推送消息   │    │    完成     │
│             │    │  (含翻译)   │    │             │    │             │    │             │
└─────────────┘    └─────────────┘    └─────────────┘    └─────────────┘    └─────────────┘
      │                  │                  │                  │
      ▼                  ▼                  ▼                  ▼
  12篇文章          5篇翻译简报        AI_News_Brief_*.md     企业微信/钉钉/飞书
  (2个源×6篇)       (自动翻译)          (Markdown格式)       (Markdown/纯文本)"""
    add_code_block(doc, flow, 'text')
    
    add_title(doc, '6.2 定时任务流程', level=2)
    flow = """启动
  │
  ▼
┌─────────────┐
│ 立即执行一次 │  ← 启动时立即爬取一次，确保用户及时获取新闻
└─────────────┘
  │
  ▼
┌───────────────────────┐
│ 计算下次运行时间       │  ← 遍历 schedule_times，找到下一个未到达的时间点
│ (09:00 / 18:00)       │     若今天已过所有时间点，返回明天第一个时间
└───────────────────────┘
  │
  ▼
┌─────────────┐
│ 等待中...   │  ← 每秒检查一次，直到到达目标时间
│ (每秒轮询)   │
└─────────────┘
  │
  ▼ (到达时间点)
┌─────────────┐
│ 执行新闻任务 │  ← 爬取新闻 → 生成简报 → 保存文件 → 推送消息
└─────────────┘
  │
  ▼ (循环)
┌─────────────┘"""
    add_code_block(doc, flow, 'text')
    
    add_title(doc, '6.3 前后端交互流程', level=2)
    flow = """用户浏览器 (http://localhost:5173)
        │
        ▼
┌───────────────────────────────────────────────────────────────┐
│  前端请求                                                      │
│  ├── GET /api/analytics/stats  → 获取统计数据                  │
│  ├── GET /api/news?page=1&limit=10 → 获取新闻列表              │
│  ├── POST /api/news → 触发爬取新闻                             │
│  ├── GET /api/config → 获取系统配置                            │
│  ├── PUT /api/config → 更新系统配置                            │
│  └── POST /api/config/push → 测试推送                          │
└───────────────────────────────────────────────────────────────┘
        │
        ▼ (HTTP请求)
┌───────────────────────────────────────────────────────────────┐
│  Flask后端 (http://localhost:5000)                             │
│  ├── 接收请求                                                  │
│  ├── 调用业务逻辑 (crawl_ai_news, generate_brief, etc.)        │
│  ├── 读写JSON文件 (articles.json, config.json, schedule.json)  │
│  └── 返回JSON响应                                              │
└───────────────────────────────────────────────────────────────┘
        │
        ▼
 用户浏览器 (渲染数据)"""
    add_code_block(doc, flow, 'text')
    
    add_title(doc, '6.4 数据流转图', level=2)
    flow = """RSS源 (TechCrunch, VentureBeat)
        │
        ▼ (爬取)
┌──────────────┐
│ crawl_ai_news│  ← 新闻爬取模块
└──────────────┘
        │
        ▼ (原始数据)
┌──────────────┐
│ extract_summary│ ← 摘要提取模块
└──────────────┘
        │
        ▼ (摘要)
┌──────────────┐
│ translate_text│ ← 翻译模块
└──────────────┘
        │
        ▼ (翻译后数据)
┌──────────────┐
│generate_brief│ ← 简报生成模块
└──────────────┘
        │
        ├───▼ (保存文件)
        │ AI_News_Brief_YYYYMMDD_HHMMSS.md
        │
        └───▼ (推送)
┌──────────────┐
│ PushNotifier │ ← 推送通知模块
└──────────────┘
        │
        ├───► 企业微信
        ├───► 钉钉
        └───► 飞书"""
    add_code_block(doc, flow, 'text')
    
    doc.add_page_break()
    
    add_title(doc, '七、技术栈分析', level=1)
    
    add_title(doc, '7.1 Python后端依赖', level=2)
    data = [
        ['依赖', '版本', '用途', '必要性'],
        ['requests', '2.34.2', 'HTTP 请求，用于爬取RSS和调用翻译API', '必须'],
        ['beautifulsoup4', '-', 'HTML/XML 解析，用于解析RSS Feed', '必须'],
        ['lxml', '6.1.1', 'XML 解析器，提高BeautifulSoup解析效率', '必须'],
        ['translators', '6.0.4', '机器翻译，支持百度/Google/腾讯翻译', '推荐'],
        ['flask', '3', 'Web框架，构建RESTful API', '必须（前端集成）'],
        ['flask-cors', '-', '跨域资源共享，支持前后端分离', '必须（前端集成）'],
        ['pyinstaller', '6.21.0', '打包工具，将Python脚本打包为exe', '构建阶段'],
        ['python-docx', '-', 'Word文档生成，用于生成分析报告', '可选'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.2 Python标准库使用', level=2)
    data = [
        ['模块', '用途', '使用位置'],
        ['time', '时间延迟、格式化', 'news_crawler.py, news_scheduler.py'],
        ['datetime', '日期时间处理、定时计算', 'news_scheduler.py, api/app.py'],
        ['re', '正则表达式、文本分割', 'simple_main.py, api/app.py'],
        ['os', '文件路径处理', '所有模块'],
        ['json', '数据序列化', 'api/app.py'],
        ['random', '随机数生成（防反爬延迟）', 'news_crawler.py'],
        ['sys', '系统路径设置', 'api/app.py, news_scheduler.py'],
        ['collections', '数据结构（defaultdict）', 'news_summary.py'],
        ['heapq', '堆排序（摘要提取）', 'news_summary.py'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.3 前端依赖', level=2)
    data = [
        ['依赖', '用途'],
        ['react', '前端框架，组件化开发'],
        ['react-dom', 'React DOM 渲染'],
        ['typescript', '类型安全，提升代码质量'],
        ['vite', '构建工具，快速热更新'],
        ['tailwindcss', 'CSS框架，快速样式开发'],
        ['postcss', 'CSS后处理器'],
        ['autoprefixer', 'CSS自动添加前缀'],
        ['chart.js', '图表库，数据可视化'],
        ['react-chartjs-2', 'Chart.js的React封装'],
        ['lucide-react', '图标库，UI图标'],
        ['axios', 'HTTP客户端，API调用'],
        ['zustand', '状态管理（轻量级）'],
        ['react-query', '数据获取和缓存'],
        ['clsx', 'CSS类名组合'],
        ['tailwind-merge', 'Tailwind CSS类名合并'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.4 架构设计特点', level=2)
    features = [
        '前后端分离：前端使用React + TypeScript，后端使用Flask，通过RESTful API通信',
        '模块化设计：各功能模块独立（爬虫、摘要、翻译、推送），便于维护和扩展',
        '松耦合：模块间通过函数调用和数据结构交互，无强依赖关系',
        '容错机制：爬取失败自动降级使用样例数据，翻译失败返回原文',
        '多层降级：翻译服务支持多种备选方案，保证翻译成功率',
        '无状态设计：每次运行独立，无持久化状态依赖（除JSON文件）',
        '配置驱动：支持通过配置文件自定义数据源、推送渠道和定时时间',
        '响应式设计：前端采用TailwindCSS，支持不同屏幕尺寸',
    ]
    add_list(doc, features)
    
    doc.add_page_break()
    
    add_title(doc, '八、部署与打包', level=1)
    
    add_title(doc, '8.1 PyInstaller 打包配置', level=2)
    add_paragraph(doc, '打包命令：', indent=True)
    code = """# 定时任务版
pyinstaller --onefile --console --name "AI_News_Crawler" news\\news_scheduler.py --clean

# 单次运行版
pyinstaller --onefile --console --name "AI_News_RunOnce" news\\run_once.py --clean

# Flask API版
pyinstaller --onefile --console --name "AI_News_API" news\\api\\app.py --clean"""
    add_code_block(doc, code)
    
    add_title(doc, '参数说明', level=3)
    data = [
        ['参数', '说明'],
        ['--onefile', '打包为单个可执行文件，便于分发'],
        ['--console', '显示控制台窗口，便于调试和查看日志'],
        ['--name', '指定输出文件名'],
        ['--clean', '清理缓存后打包，避免旧版本干扰'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '8.2 打包产物结构', level=2)
    code = """dist/
├── AI_News_Crawler.exe        # 定时任务版可执行文件 (~50MB)
│   ├── Python 解释器
│   ├── 所有依赖库 (requests, beautifulsoup4, etc.)
│   ├── news_scheduler.py
│   ├── simple_main.py
│   ├── push_notifier.py
│   └── 资源文件
├── AI_News_RunOnce.exe        # 单次运行版可执行文件 (~50MB)
├── AI_News_API.exe            # Flask API版可执行文件 (~50MB)
├── config.json                # 系统配置文件
├── setup_task_scheduler.ps1   # Windows任务计划配置脚本
└── start_news.bat             # 启动脚本"""
    add_code_block(doc, code)
    
    add_title(doc, '8.3 Windows 任务计划配置', level=2)
    add_paragraph(doc, '推荐配置：', indent=True)
    data = [
        ['配置项', '值'],
        ['任务名称', 'AI_News_Crawler'],
        ['触发器', '开机时 + 每天 09:00 + 每天 18:00'],
        ['操作', '启动程序'],
        ['程序路径', 'D:\\ruangong233\\demo2\\dist\\AI_News_RunOnce.exe'],
        ['工作目录', 'D:\\ruangong233\\demo2\\dist'],
        ['安全选项', '不管用户是否登录都要运行'],
        ['安全选项', '以最高权限运行'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '8.4 前端部署', level=2)
    code = """# 开发环境启动
cd frontend
npm install
npm run dev

# 生产环境构建
cd frontend
npm install
npm run build

# 构建产物
frontend/dist/
├── index.html
├── assets/
│   ├── index-xxx.js
│   ├── index-xxx.css
│   └── vendor-xxx.js
└── favicon.svg

# 生产环境部署（使用Nginx）
server {
    listen 80;
    server_name ai-news.example.com;
    
    location / {
        root /path/to/frontend/dist;
        index index.html;
        try_files $uri $uri/ /index.html;
    }
    
    location /api/ {
        proxy_pass http://localhost:5000/api/;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
    }
}"""
    add_code_block(doc, code)
    
    doc.add_page_break()
    
    add_title(doc, '九、系统运行分析', level=1)
    
    add_title(doc, '9.1 执行时间分析', level=2)
    data = [
        ['步骤', '平均耗时', '影响因素'],
        ['爬取新闻', '5-10秒', '网络状况、目标网站响应速度、防反爬延迟'],
        ['摘要提取', '<1秒', '文本长度、句子数量'],
        ['翻译处理', '3-5秒/篇', '翻译服务响应速度、网络状况'],
        ['生成简报', '<1秒', '本地计算'],
        ['保存文件', '<1秒', '磁盘IO速度'],
        ['推送消息', '1-3秒/平台', '网络状况、平台响应速度'],
        ['总计（单次）', '15-30秒', '网络状况、新闻数量'],
        ['定时任务间隔', '9小时或12小时', '配置的运行时间'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '9.2 资源消耗', level=2)
    data = [
        ['指标', '运行时峰值', '打包后大小'],
        ['内存', '~100MB', '~50MB (exe)'],
        ['CPU', '中等（爬取时）', '-'],
        ['网络', '约 100KB/次', '-'],
        ['磁盘', '约 500KB/简报文件', '-'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '9.3 错误处理机制', level=2)
    data = [
        ['异常场景', '处理方式', '影响'],
        ['网络请求失败', '使用预设样例数据', '系统继续运行，数据为样例数据'],
        ['翻译服务不可用', '返回原文', '简报为英文，不影响推送'],
        ['推送失败', '记录错误日志，继续执行', '消息未推送，但简报已保存'],
        ['文件写入失败', '打印错误信息', '简报未保存，但推送正常'],
        ['键盘中断', '优雅退出', '程序正常结束'],
        ['API请求异常', '返回错误状态码和消息', '前端显示错误提示'],
        ['数据文件损坏', '返回默认配置或空数组', '系统使用默认配置或空数据'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '十、优缺点评估', level=1)
    
    add_title(doc, '10.1 优点', level=2)
    data = [
        ['类别', '描述'],
        ['功能完整', '涵盖采集、摘要、翻译、推送、前端分析全流程'],
        ['易于部署', '打包为单文件，无需安装 Python 环境'],
        ['多平台推送', '支持企业微信、钉钉、飞书三大主流平台'],
        ['容错性强', '多层降级策略，保证系统可用性'],
        ['无外部依赖', '核心功能使用标准库实现，降低部署难度'],
        ['可配置性', '支持自定义推送渠道、数据源和定时时间'],
        ['前后端分离', 'RESTful API 设计，支持多端调用'],
        ['数据可视化', '前端图表展示新闻趋势、来源分布、关键词'],
        ['响应式界面', '深蓝色主题，支持移动端和桌面端'],
        ['代码质量', '模块化设计，代码结构清晰，易于维护'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '10.2 缺点与改进建议', level=2)
    data = [
        ['类别', '问题', '改进建议'],
        ['翻译质量', '依赖免费翻译API，质量有限', '接入商用翻译服务（如百度翻译API），提高翻译准确率'],
        ['摘要算法', '基于词频统计，缺乏语义理解', '引入 TF-IDF 或 BERT 模型，实现语义级摘要'],
        ['数据源', '仅支持 2 个 RSS 源', '添加更多数据源（如 MIT Tech Review、ZDNet），支持自定义源'],
        ['数据持久化', 'JSON文件存储，不支持复杂查询', '引入 SQLite 或 MongoDB，支持更复杂的数据查询和分析'],
        ['并发处理', '单线程顺序执行，效率较低', '引入多线程/异步处理，提高爬取和翻译效率'],
        ['配置管理', '部分配置硬编码在代码中', '完善配置文件，支持环境变量和热更新'],
        ['日志系统', '仅控制台输出，无文件日志', '添加文件日志和日志级别控制，便于问题排查'],
        ['错误告警', '无告警机制', '添加异常邮件/消息告警，及时通知运维人员'],
        ['前端状态管理', '使用简单的 useState', '引入 Zustand 或 Redux，统一管理全局状态'],
        ['数据缓存', '前端无缓存机制', '使用 React Query 或 SWR，实现数据缓存和自动更新'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '十一、扩展建议', level=1)
    
    add_title(doc, '11.1 功能扩展', level=2)
    features = [
        '多语言支持：增加日语、韩语、法语等其他语言翻译',
        '图片采集：从新闻中提取配图，丰富简报内容',
        '视频新闻：支持视频内容采集和链接提取',
        '智能分类：基于机器学习的新闻自动分类（技术、商业、伦理等）',
        '情感分析：分析新闻情感倾向（正面、负面、中性）',
        '关键词追踪：自定义关键词订阅，实时推送相关新闻',
        '个性化推荐：基于用户阅读历史的智能推荐',
        '数据导出：支持将新闻数据导出为 Excel、CSV 格式',
        '简报模板：支持自定义简报模板和样式',
        '历史对比：支持不同时间段新闻数据的对比分析',
    ]
    add_list(doc, features)
    
    add_title(doc, '11.2 架构优化', level=2)
    optimizations = [
        '微服务架构：将爬虫、翻译、推送拆分为独立服务',
        '消息队列：使用 Redis/Kafka 实现异步处理和解耦',
        '缓存机制：缓存已爬取的新闻，避免重复采集',
        '分布式部署：支持多节点分布式爬取，提高效率',
        '容器化：使用 Docker 容器化部署，便于环境管理',
        '监控系统：添加 Prometheus/Grafana 监控，实时监控系统状态',
        '负载均衡：添加 Nginx 负载均衡，支持高并发访问',
        '熔断降级：使用 Hystrix 实现服务熔断，防止级联故障',
    ]
    add_list(doc, optimizations)
    
    add_title(doc, '11.3 运维优化', level=2)
    ops = [
        '健康检查：定期检查服务状态，自动发现问题',
        '自动重启：服务异常时自动重启，提高可用性',
        '版本管理：支持版本自动更新和回滚',
        '配置中心：集中管理配置，支持动态配置',
        '灰度发布：支持灰度发布策略，降低发布风险',
        '日志分析：使用 ELK 栈进行日志收集和分析',
        '性能监控：实时监控系统性能指标',
        '安全加固：添加 API 认证、请求限流、IP 白名单',
    ]
    add_list(doc, ops)
    
    doc.add_page_break()
    
    add_title(doc, '十二、总结', level=1)
    
    add_title(doc, '12.1 系统价值', level=2)
    values = [
        '信息获取效率提升：自动化采集，节省人工浏览时间',
        '信息质量保证：摘要提取和翻译处理，提供结构化内容',
        '多渠道触达：支持多种即时通讯工具推送，随时随地获取资讯',
        '部署便捷：单文件打包，零依赖运行，适合个人和企业使用',
        '实时动态分析：前端可视化展示，直观了解新闻趋势和热点',
        '可扩展性强：模块化设计，便于后续功能扩展和维护',
    ]
    add_list(doc, values)
    
    add_title(doc, '12.2 适用场景', level=2)
    scenarios = [
        '个人使用：每日获取 AI 新闻简报，了解行业动态',
        '团队协作：团队内部共享技术资讯，促进知识交流',
        '企业应用：企业内部分发行业动态，辅助决策支持',
        '学习研究：AI 研究者跟踪最新进展，获取研究资料',
        '媒体监控：品牌或产品相关新闻监控，及时了解舆情',
    ]
    add_list(doc, scenarios)
    
    add_title(doc, '12.3 技术贡献', level=2)
    contributions = [
        '实现了基于词频统计的文本摘要算法，无需外部依赖',
        '设计了多层降级翻译策略，提高翻译成功率',
        '开发了多平台消息推送模块，统一接口封装',
        '实现了基于标准库的定时任务调度器',
        '构建了完整的前后端分离架构，RESTful API 设计',
        '开发了数据可视化前端界面，支持实时动态分析',
    ]
    add_list(doc, contributions)
    
    add_title(doc, '12.4 未来展望', level=2)
    add_paragraph(doc, '随着 AI 技术的快速发展，该系统可以进一步集成更多先进技术，提升智能化水平：', indent=True)
    future = [
        '大语言模型集成：使用 GPT、Claude 等大语言模型进行更智能的摘要、分析和生成',
        '多模态内容：支持图文、视频等多种内容形式的采集和展示',
        '个性化推荐：基于用户偏好和阅读习惯的智能推荐系统',
        '实时推送：支持实时新闻推送，不再依赖定时任务，即时获取最新资讯',
        '知识图谱：构建 AI 领域知识图谱，实现更深入的关联分析',
        '智能问答：基于新闻数据的智能问答系统，支持自然语言查询',
    ]
    add_list(doc, future)
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('— 报告结束 —')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = os.path.join(os.path.dirname(__file__), 'AI_新闻爬虫系统完整分析报告.docx')
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_report()