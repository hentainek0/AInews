from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

def add_title(doc, text, level=1):
    if level == 1:
        style = 'Heading 1'
    elif level == 2:
        style = 'Heading 2'
    elif level == 3:
        style = 'Heading 3'
    else:
        style = 'Heading 4'
    
    title = doc.add_heading(text, level=level)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = '微软雅黑'
    title.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return title

def add_paragraph(doc, text, font_size=10.5, bold=False, indent=False, monospace=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    if monospace:
        run.font.name = 'Consolas'
    else:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, data, headers=None):
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if headers and i == 0:
                cell.text = headers[j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                cell.text = data[i - 1 if headers else i][j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    table.style = 'Table Grid'
    return table

def add_list(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        if ordered:
            p = doc.add_paragraph(f"{i}. {item}", style='List Number')
        else:
            p = doc.add_paragraph(f"• {item}", style='List Bullet')
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def generate_usecase_doc():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    add_title(doc, 'AI新闻爬虫系统', level=1)
    add_title(doc, '用例图与用例描述表', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(f'版本：v1.0')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(f'编制日期：{datetime.datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    add_title(doc, '目录', level=1)
    toc_items = [
        ('第一章 系统概述', '1'),
        ('第二章 用例图', '3'),
        ('第三章 用例描述表', '8'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}'.ljust(40) + f'第 {page} 页')
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    add_title(doc, '第一章 系统概述', level=1)
    
    add_title(doc, '1.1 系统简介', level=2)
    add_paragraph(doc, 'AI 新闻爬虫系统是一个自动化的新闻采集、分析与推送平台，主要功能包括：从多个 RSS 数据源采集 AI 相关新闻，对新闻内容进行摘要提取和翻译处理，通过企业微信、钉钉、飞书等平台推送新闻简报，提供 Web 前端界面进行实时动态分析，支持定时任务自动执行新闻采集和推送。', indent=True)
    
    add_title(doc, '1.2 参与者定义', level=2)
    data = [
        ['参与者', '编号', '描述', '职责'],
        ['系统管理员', 'Actor-001', '负责系统配置和维护的人员', '配置数据源、推送渠道、定时任务，管理系统运行'],
        ['AI 从业者', 'Actor-002', 'AI 研发工程师、产品经理等', '浏览新闻、查看分析数据、获取技术资讯'],
        ['企业团队', 'Actor-003', '科技公司技术团队', '接收新闻简报、共享技术资讯'],
        ['系统', 'Actor-004', '系统自身', '自动执行定时任务、爬取新闻、处理文本、推送消息'],
        ['外部服务', 'Actor-005', '第三方 API 服务', '提供翻译服务、推送服务、RSS 数据源'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '1.3 系统边界', level=2)
    add_paragraph(doc, '系统边界定义了 AI 新闻爬虫系统与外部参与者之间的交互范围：', indent=True)
    
    boundaries = [
        '系统内部：新闻爬虫模块、文本处理模块、推送模块、定时任务模块、后端 API、前端界面、数据存储',
        '系统外部参与者：系统管理员、AI 从业者、企业团队',
        '外部服务接口：企业微信机器人 API、钉钉机器人 API、飞书机器人 API、翻译服务 API、RSS 数据源',
    ]
    add_list(doc, boundaries)
    
    doc.add_page_break()
    
    add_title(doc, '第二章 用例图', level=1)
    
    add_title(doc, '2.1 用例图概述', level=2)
    add_paragraph(doc, '用例图是 UML 中用于描述系统功能的图形化表示，展示了系统的参与者、用例以及它们之间的关系。本系统的用例图按功能模块分为六个部分：新闻爬虫模块、文本处理模块、推送模块、定时任务模块、前端界面模块和后端 API 模块。', indent=True)
    
    add_title(doc, '2.2 用例图 PlantUML 描述', level=2)
    add_paragraph(doc, '以下使用 PlantUML 语法描述系统的完整用例图：', indent=True)
    
    plantuml_code = """@startuml AI新闻爬虫系统用例图
skinparam backgroundColor #FEFEFE
skinparam handwritten false

actor "系统管理员" as Admin #90EE90
actor "AI从业者" as User #87CEEB
actor "企业团队" as Team #DDA0DD
actor "系统" as System #FFD700
actor "外部服务" as External #F0E68C

rectangle "AI新闻爬虫系统" {
  package "新闻爬虫模块" {
    usecase "定时爬取新闻" as UC001
    usecase "手动爬取新闻" as UC002
    usecase "解析RSS内容" as UC007
    usecase "提取新闻信息" as UC008
  }
  
  package "文本处理模块" {
    usecase "提取新闻摘要" as UC003
    usecase "翻译新闻内容" as UC004
    usecase "提取关键词标签" as UC009
  }
  
  package "推送模块" {
    usecase "推送新闻简报" as UC005
    usecase "测试推送配置" as UC010
    usecase "记录推送结果" as UC011
  }
  
  package "定时任务模块" {
    usecase "定时执行任务" as UC006
    usecase "配置执行时间" as UC012
    usecase "记录任务日志" as UC013
  }
  
  package "前端界面模块" {
    usecase "查看Dashboard" as UC014
    usecase "浏览新闻列表" as UC015
    usecase "查看新闻详情" as UC016
    usecase "查看数据分析" as UC017
    usecase "配置系统设置" as UC018
  }
  
  package "后端API模块" {
    usecase "管理新闻数据" as UC019
    usecase "获取统计数据" as UC020
    usecase "管理系统配置" as UC021
    usecase "触发手动推送" as UC022
  }
}

Admin --> UC002 : 手动触发
Admin --> UC010 : 测试配置
Admin --> UC012 : 配置时间
Admin --> UC018 : 配置设置
Admin --> UC021 : 管理配置
Admin --> UC022 : 手动推送

User --> UC014 : 查看概览
User --> UC015 : 浏览新闻
User --> UC016 : 查看详情
User --> UC017 : 分析数据

Team --> UC005 : 接收简报

System --> UC001 : 自动执行
System --> UC003 : 自动处理
System --> UC004 : 自动翻译
System --> UC006 : 定时触发
System --> UC007 : 解析内容
System --> UC008 : 提取信息
System --> UC009 : 提取标签
System --> UC011 : 记录结果
System --> UC013 : 记录日志
System --> UC019 : 管理数据
System --> UC020 : 获取统计

External --> UC004 : 提供翻译
External --> UC005 : 发送消息
External --> UC007 : 提供数据

UC001 --> UC007 : include
UC001 --> UC008 : include
UC001 --> UC003 : include
UC001 --> UC004 : include
UC001 --> UC005 : include

UC002 --> UC007 : include
UC002 --> UC008 : include
UC002 --> UC003 : include
UC002 --> UC004 : include

UC006 --> UC001 : include

UC005 --> UC011 : include

@enduml"""
    
    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run(plantuml_code)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'
    
    add_title(doc, '2.3 用例关系说明', level=2)
    add_paragraph(doc, '用例图中的关系包括以下几种类型：', indent=True)
    
    relationships = [
        {
            '类型': '关联关系 (Association)',
            '符号': '实线箭头',
            '说明': '表示参与者与用例之间的交互关系，即参与者使用或触发用例',
            '示例': '系统管理员 → 手动爬取新闻，AI从业者 → 浏览新闻列表',
        },
        {
            '类型': '包含关系 (Include)',
            '符号': '虚线箭头 <<include>>',
            '说明': '表示一个用例包含另一个用例的行为，被包含的用例是基用例的一部分',
            '示例': '定时爬取新闻 include 解析RSS内容、提取新闻信息、提取新闻摘要、翻译新闻内容、推送新闻简报',
        },
        {
            '类型': '扩展关系 (Extend)',
            '符号': '虚线箭头 <<extend>>',
            '说明': '表示一个用例可以扩展另一个用例的行为，扩展用例在特定条件下执行',
            '示例': '（本系统暂未使用扩展关系）',
        },
        {
            '类型': '泛化关系 (Generalization)',
            '符号': '实线空心三角箭头',
            '说明': '表示一个用例是另一个用例的特殊化，子类用例继承父类用例的行为',
            '示例': '（本系统暂未使用泛化关系）',
        },
    ]
    
    for rel in relationships:
        add_paragraph(doc, f"**{rel['类型']}**", bold=True)
        add_paragraph(doc, f"符号：{rel['符号']}")
        add_paragraph(doc, f"说明：{rel['说明']}")
        add_paragraph(doc, f"示例：{rel['示例']}")
        doc.add_paragraph()
    
    add_title(doc, '2.4 模块用例图', level=2)
    
    modules = [
        {
            'name': '新闻爬虫模块',
            'description': '负责从多个 RSS 数据源采集新闻',
            'usecases': ['UC-001 定时爬取新闻', 'UC-002 手动爬取新闻', 'UC-007 解析RSS内容', 'UC-008 提取新闻信息'],
            'actors': ['系统管理员', '系统', '外部服务'],
        },
        {
            'name': '文本处理模块',
            'description': '负责新闻摘要提取和翻译处理',
            'usecases': ['UC-003 提取新闻摘要', 'UC-004 翻译新闻内容', 'UC-009 提取关键词标签'],
            'actors': ['系统', '外部服务'],
        },
        {
            'name': '推送模块',
            'description': '负责新闻简报推送',
            'usecases': ['UC-005 推送新闻简报', 'UC-010 测试推送配置', 'UC-011 记录推送结果'],
            'actors': ['系统管理员', '企业团队', '系统', '外部服务'],
        },
        {
            'name': '定时任务模块',
            'description': '负责定时执行新闻采集和推送任务',
            'usecases': ['UC-006 定时执行任务', 'UC-012 配置执行时间', 'UC-013 记录任务日志'],
            'actors': ['系统管理员', '系统'],
        },
        {
            'name': '前端界面模块',
            'description': '提供用户交互界面',
            'usecases': ['UC-014 查看Dashboard', 'UC-015 浏览新闻列表', 'UC-016 查看新闻详情', 'UC-017 查看数据分析', 'UC-018 配置系统设置'],
            'actors': ['系统管理员', 'AI从业者'],
        },
        {
            'name': '后端API模块',
            'description': '提供 RESTful API 服务',
            'usecases': ['UC-019 管理新闻数据', 'UC-020 获取统计数据', 'UC-021 管理系统配置', 'UC-022 触发手动推送'],
            'actors': ['系统管理员', '系统'],
        },
    ]
    
    for module in modules:
        add_title(doc, f"2.4.{modules.index(module) + 1} {module['name']}", level=3)
        add_paragraph(doc, module['description'])
        
        add_paragraph(doc, "**用例列表**：")
        add_list(doc, module['usecases'])
        
        add_paragraph(doc, "**参与者**：")
        add_list(doc, module['actors'])
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    add_title(doc, '第三章 用例描述表', level=1)
    
    add_title(doc, '3.1 用例描述表说明', level=2)
    add_paragraph(doc, '用例描述表详细描述了每个用例的属性和流程，包括用例编号、名称、描述、参与者、前置条件、后置条件、基本流程、备选流程、异常流程和优先级等字段。', indent=True)
    
    add_title(doc, '3.2 新闻爬虫模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-001', '定时爬取新闻', '系统按照预定时间自动爬取新闻', '系统', '系统已启动，定时任务已配置', '新闻数据已更新到数据库', '1.系统到达预定时间\n2.系统调用爬虫接口\n3.爬虫请求RSS Feed\n4.解析RSS内容\n5.提取新闻标题、链接、内容、来源\n6.保存新闻数据', '无', '1.网络异常：重试3次后记录错误日志\n2.RSS Feed格式错误：跳过该数据源', '高'],
        ['UC-002', '手动爬取新闻', '用户手动触发新闻爬取', '系统管理员', '用户已登录系统', '新闻数据已更新，用户收到爬取结果通知', '1.用户点击手动爬取按钮\n2.系统调用爬虫接口\n3.爬虫请求RSS Feed\n4.解析RSS内容\n5.提取新闻信息\n6.保存新闻数据\n7.返回爬取结果', '无', '1.网络异常：显示错误提示\n2.爬虫服务未启动：提示用户启动服务', '高'],
        ['UC-007', '解析RSS内容', '解析RSS Feed返回的XML内容', '系统', '已获取RSS Feed内容', '返回解析后的新闻列表', '1.接收RSS Feed内容\n2.解析XML结构\n3.提取新闻条目\n4.返回解析结果', '无', '1.XML格式错误：返回空列表\n2.编码错误：尝试多种编码解析', '高'],
        ['UC-008', '提取新闻信息', '从新闻条目中提取标题、链接、内容等信息', '系统', '已解析RSS内容', '返回提取后的新闻信息', '1.遍历新闻条目\n2.提取标题\n3.提取链接\n4.提取发布日期\n5.提取内容摘要\n6.提取来源信息\n7.返回新闻信息列表', '无', '1.字段缺失：使用默认值\n2.内容过长：截取前200字', '高'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.3 文本处理模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-003', '提取新闻摘要', '系统对新闻内容进行摘要提取', '系统', '新闻内容已获取', '摘要已保存到新闻数据中', '1.系统接收新闻内容\n2.使用摘要算法提取关键信息\n3.生成200字以内的摘要\n4.保存摘要内容', '无', '1.新闻内容过短：直接使用原文作为摘要', '高'],
        ['UC-004', '翻译新闻内容', '系统将英文新闻翻译为中文', '系统', '新闻内容已获取', '翻译结果已保存到新闻数据中', '1.系统检测新闻语言\n2.若为英文，调用翻译服务\n3.翻译标题和内容\n4.保存翻译结果\n5.翻译失败时使用原文', '无', '1.翻译服务不可用：降级使用备用翻译服务\n2.翻译超时：使用原文', '高'],
        ['UC-009', '提取关键词标签', '从新闻标题和内容中提取关键词', '系统', '新闻内容已获取', '关键词标签已保存到新闻数据中', '1.系统接收新闻标题和内容\n2.使用正则表达式提取英文单词\n3.过滤停用词\n4.统计词频\n5.选取Top10关键词作为标签\n6.保存标签', '无', '1.关键词过少：使用默认标签\n2.无英文内容：提取中文关键词', '中'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.4 推送模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-005', '推送新闻简报', '系统将新闻简报推送到配置的渠道', '系统、企业团队', '新闻简报已生成，推送配置已完成', '推送结果已记录', '1.系统获取推送配置\n2.生成推送内容\n3.调用企业微信接口推送\n4.调用钉钉接口推送\n5.调用飞书接口推送\n6.记录推送结果', '无', '1.推送失败：记录错误日志，不影响其他渠道', '高'],
        ['UC-010', '测试推送配置', '测试推送配置是否正确', '系统管理员', '推送配置已填写', '用户收到测试消息', '1.用户点击测试按钮\n2.系统生成测试消息\n3.调用推送接口发送消息\n4.返回测试结果', '无', '1.配置错误：显示错误信息\n2.网络异常：显示错误信息', '中'],
        ['UC-011', '记录推送结果', '记录每次推送的结果', '系统', '推送任务已执行', '推送结果已保存到日志', '1.接收推送结果\n2.记录推送时间\n3.记录推送渠道\n4.记录推送状态\n5.记录错误信息（如有）\n6.保存日志', '无', '1.日志写入失败：尝试备用存储', '中'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.5 定时任务模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-006', '定时执行任务', '系统按照配置的时间自动执行任务', '系统', '定时任务已配置，系统已启动', '任务执行完成，结果已记录', '1.系统检测当前时间\n2.匹配预定时间点\n3.执行新闻爬取任务\n4.执行新闻摘要任务\n5.执行新闻翻译任务\n6.执行新闻推送任务\n7.记录任务执行结果', '无', '1.任务执行失败：记录错误日志，下次继续执行', '高'],
        ['UC-012', '配置执行时间', '配置定时任务的执行时间', '系统管理员', '用户已登录系统', '执行时间已保存', '1.用户进入定时任务设置\n2.输入执行时间\n3.保存配置\n4.系统更新定时任务', '无', '1.时间格式错误：提示重新输入', '高'],
        ['UC-013', '记录任务日志', '记录定时任务的执行日志', '系统', '定时任务已执行', '日志已保存', '1.记录任务开始时间\n2.记录任务结束时间\n3.记录任务状态\n4.记录执行结果\n5.记录错误信息（如有）\n6.保存日志', '无', '1.日志写入失败：尝试备用存储', '中'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.6 前端界面模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-014', '查看Dashboard', '查看系统概览数据', '系统管理员、AI从业者', '用户已打开前端页面', 'Dashboard显示最新数据', '1.用户打开Dashboard页面\n2.系统获取统计数据\n3.展示新闻总数、今日新增\n4.展示近7天趋势图表\n5.展示来源分布\n6.展示关键词统计', '无', '1.数据获取失败：显示加载错误', '高'],
        ['UC-015', '浏览新闻列表', '查看新闻列表', '系统管理员、AI从业者', '用户已打开前端页面', '新闻列表显示', '1.用户打开新闻列表页面\n2.系统获取新闻数据\n3.分页展示新闻列表\n4.支持按来源筛选\n5.支持关键词搜索', '无', '1.数据获取失败：显示加载错误', '高'],
        ['UC-016', '查看新闻详情', '查看新闻详细内容', '系统管理员、AI从业者', '用户已选择新闻', '新闻详情显示', '1.用户点击新闻标题\n2.系统获取新闻详情\n3.展示新闻标题\n4.展示新闻内容\n5.展示新闻链接\n6.展示发布时间和来源', '无', '1.新闻不存在：显示404错误', '高'],
        ['UC-017', '查看数据分析', '查看新闻数据分析', '系统管理员、AI从业者', '用户已打开前端页面', '数据分析展示', '1.用户打开数据分析页面\n2.系统获取分析数据\n3.展示来源分布图表\n4.展示关键词统计图表\n5.展示趋势分析图表', '无', '1.数据获取失败：显示加载错误', '高'],
        ['UC-018', '配置系统设置', '配置系统参数', '系统管理员', '用户已登录系统', '配置已保存', '1.用户进入系统设置页面\n2.配置推送渠道\n3.配置定时任务\n4.配置数据源\n5.保存配置', '无', '1.配置格式错误：提示重新输入', '高'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.7 后端API模块用例', level=2)
    
    data = [
        ['用例编号', '用例名称', '描述', '参与者', '前置条件', '后置条件', '基本流程', '备选流程', '异常流程', '优先级'],
        ['UC-019', '管理新闻数据', '管理新闻数据（增删查）', '系统管理员、系统', '系统已启动', '新闻数据已更新', '1.接收API请求\n2.验证请求参数\n3.执行增删查操作\n4.返回操作结果', '无', '1.参数错误：返回错误信息\n2.操作失败：返回错误信息', '高'],
        ['UC-020', '获取统计数据', '获取新闻统计数据', '系统管理员、系统', '系统已启动', '统计数据已返回', '1.接收API请求\n2.查询新闻数据\n3.计算统计指标\n4.返回统计结果', '无', '1.数据为空：返回默认值', '高'],
        ['UC-021', '管理系统配置', '管理系统配置信息', '系统管理员', '用户已登录系统', '配置已更新', '1.接收配置更新请求\n2.验证配置参数\n3.保存配置信息\n4.返回操作结果', '无', '1.配置格式错误：返回错误信息', '高'],
        ['UC-022', '触发手动推送', '手动触发新闻推送', '系统管理员', '推送配置已完成', '推送任务已执行', '1.接收推送请求\n2.生成新闻简报\n3.调用推送服务\n4.返回推送结果', '无', '1.推送失败：返回错误信息', '中'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.8 用例优先级汇总', level=2)
    data = [
        ['优先级', '用例编号', '用例名称', '数量'],
        ['高', 'UC-001', '定时爬取新闻', '13'],
        ['高', 'UC-002', '手动爬取新闻', '-'],
        ['高', 'UC-003', '提取新闻摘要', '-'],
        ['高', 'UC-004', '翻译新闻内容', '-'],
        ['高', 'UC-005', '推送新闻简报', '-'],
        ['高', 'UC-006', '定时执行任务', '-'],
        ['高', 'UC-007', '解析RSS内容', '-'],
        ['高', 'UC-008', '提取新闻信息', '-'],
        ['高', 'UC-012', '配置执行时间', '-'],
        ['高', 'UC-014', '查看Dashboard', '-'],
        ['高', 'UC-015', '浏览新闻列表', '-'],
        ['高', 'UC-016', '查看新闻详情', '-'],
        ['高', 'UC-017', '查看数据分析', '-'],
        ['高', 'UC-018', '配置系统设置', '-'],
        ['高', 'UC-019', '管理新闻数据', '-'],
        ['高', 'UC-020', '获取统计数据', '-'],
        ['高', 'UC-021', '管理系统配置', '-'],
        ['中', 'UC-009', '提取关键词标签', '6'],
        ['中', 'UC-010', '测试推送配置', '-'],
        ['中', 'UC-011', '记录推送结果', '-'],
        ['中', 'UC-013', '记录任务日志', '-'],
        ['中', 'UC-022', '触发手动推送', '-'],
        ['总计', '-', '-', '19'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('— 文档结束 —')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = os.path.join(os.path.dirname(__file__), 'AI新闻爬虫系统用例图与用例描述表.docx')
    doc.save(output_path)
    print(f"用例图与用例描述表已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_usecase_doc()
