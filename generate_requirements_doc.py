from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

def add_title(doc, text, level=1):
    if level == 1:
        style = 'Heading 1'
    elif level == 2:
        style = 'Heading 2'
    elif level == 3:
        style = 'Heading 3'
    else:
        style = 'Heading 4'
    
    title = doc.add_heading(text, level=level)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = '微软雅黑'
    title.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return title

def add_paragraph(doc, text, font_size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, data, headers=None):
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if headers and i == 0:
                cell.text = headers[j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                cell.text = data[i - 1 if headers else i][j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    table.style = 'Table Grid'
    return table

def add_list(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        if ordered:
            p = doc.add_paragraph(f"{i}. {item}", style='List Number')
        else:
            p = doc.add_paragraph(f"• {item}", style='List Bullet')
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_usecase(doc, id, name, description, actor, precondition, flow, postcondition, exceptions):
    add_paragraph(doc, f"**用例编号**：{id}", bold=True)
    add_paragraph(doc, f"**用例名称**：{name}", bold=True)
    add_paragraph(doc, f"**用例描述**：{description}")
    add_paragraph(doc, f"**参与者**：{actor}")
    add_paragraph(doc, f"**前置条件**：{precondition}")
    add_paragraph(doc, "**基本流程**：")
    for i, step in enumerate(flow, 1):
        add_paragraph(doc, f"{i}. {step}", indent=True)
    add_paragraph(doc, f"**后置条件**：{postcondition}")
    add_paragraph(doc, "**异常流程**：")
    for i, exc in enumerate(exceptions, 1):
        add_paragraph(doc, f"{i}. {exc}", indent=True)

def generate_requirements_doc():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    add_title(doc, 'AI新闻爬虫系统', level=1)
    add_title(doc, '需求规格说明书', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(f'版本：v1.0')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(f'编制日期：{datetime.datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    add_title(doc, '目录', level=1)
    toc_items = [
        ('第一章 引言', '1'),
        ('第二章 项目概述', '3'),
        ('第三章 用户角色与场景分析', '8'),
        ('第四章 功能需求', '12'),
        ('第五章 非功能需求', '35'),
        ('第六章 数据需求', '42'),
        ('第七章 接口需求', '48'),
        ('第八章 约束条件', '55'),
        ('第九章 验收标准', '58'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}'.ljust(40) + f'第 {page} 页')
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    add_title(doc, '第一章 引言', level=1)
    
    add_title(doc, '1.1 文档目的', level=2)
    add_paragraph(doc, '本文档是 AI 新闻爬虫系统的需求规格说明书（SRS），旨在详细描述系统的功能需求、非功能需求、数据需求和接口需求，为开发团队提供明确的开发依据，同时为测试团队提供验收标准。本文档适用于项目开发人员、测试人员、产品经理和项目管理者。', indent=True)
    
    add_title(doc, '1.2 术语定义', level=2)
    data = [
        ['术语', '英文', '定义'],
        ['RSS', 'Really Simple Syndication', '简易信息聚合，一种用于发布经常更新的内容的格式'],
        ['摘要', 'Summary', '对新闻内容的精炼概括，保留核心信息'],
        ['翻译', 'Translation', '将一种语言转换为另一种语言'],
        ['Webhook', '-', '一种 HTTP 回调，用于实时接收事件通知'],
        ['推送', 'Push', '将消息主动发送到用户终端'],
        ['定时任务', 'Scheduled Task', '按预定时间自动执行的任务'],
        ['RESTful API', '-', '基于 REST 架构风格的应用程序接口'],
        ['CORS', 'Cross-Origin Resource Sharing', '跨域资源共享'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '1.3 参考资料', level=2)
    refs = [
        '《AI新闻爬虫系统可行性分析与开发计划报告》',
        '《企业微信机器人接口文档》',
        '《钉钉自定义机器人开发文档》',
        '《飞书开放平台文档》',
        'Flask 官方文档：https://flask.palletsprojects.com/',
        'React 官方文档：https://react.dev/',
        'TailwindCSS 官方文档：https://tailwindcss.com/',
    ]
    add_list(doc, refs)
    
    doc.add_page_break()
    
    add_title(doc, '第二章 项目概述', level=1)
    
    add_title(doc, '2.1 项目背景', level=2)
    add_paragraph(doc, '随着人工智能技术的飞速发展，AI 领域的新闻资讯呈指数级增长。企业、团队和个人需要及时了解 AI 技术的最新动态、行业趋势和市场变化。然而，人工浏览多个科技媒体网站效率低下，且容易遗漏重要信息。', indent=True)
    
    add_paragraph(doc, '本项目旨在开发一套自动化的 AI 新闻采集、分析与推送系统，帮助用户高效获取 AI 领域的最新资讯，提升信息获取效率和决策质量。', indent=True)
    
    add_title(doc, '2.2 项目目标', level=2)
    add_paragraph(doc, '本项目的核心目标是构建一个完整的 AI 新闻自动化处理平台，具体包括：', indent=True)
    goals = [
        '自动化采集：从国际主流科技媒体自动抓取 AI 相关新闻',
        '智能分析：对新闻内容进行摘要提取和翻译处理',
        '多端推送：通过企业微信、钉钉、飞书等平台推送新闻简报',
        '实时监控：提供 Web 前端界面进行实时动态分析',
        '定时运行：支持定时任务自动执行新闻采集和推送',
        '数据管理：支持新闻数据的存储、查询和导出',
    ]
    add_list(doc, goals, ordered=True)
    
    add_title(doc, '2.3 系统范围', level=2)
    add_paragraph(doc, '本系统的功能范围包括以下几个方面：', indent=True)
    
    scope_in = [
        '新闻爬虫模块：从多个 RSS 数据源采集 AI 相关新闻',
        '文本处理模块：新闻摘要提取、中英文翻译',
        '推送模块：企业微信、钉钉、飞书消息推送',
        '定时任务模块：定时执行新闻采集和推送',
        '后端 API：提供新闻管理、分析和配置接口',
        '前端界面：Dashboard、新闻列表、数据分析、系统设置',
        '数据存储：新闻数据、系统配置的持久化存储',
    ]
    add_paragraph(doc, '**包含范围**：', bold=True)
    add_list(doc, scope_in)
    
    scope_out = [
        '新闻内容的原创创作',
        '用户注册和登录系统',
        '用户权限管理',
        '多租户支持',
        '移动端 APP 开发',
        '商业化运营功能（如广告、付费订阅）',
        '视频内容的爬取和处理',
    ]
    add_paragraph(doc, '**不包含范围**：', bold=True)
    add_list(doc, scope_out)
    
    add_title(doc, '2.4 系统架构', level=2)
    add_paragraph(doc, '本系统采用前后端分离的架构设计：', indent=True)
    
    arch = [
        '前端层：React + TypeScript + TailwindCSS，提供用户交互界面',
        '后端层：Flask + Python，提供 RESTful API 服务',
        '数据层：JSON 文件 / SQLite，存储新闻数据和配置信息',
        '服务层：爬虫服务、翻译服务、推送服务、定时任务服务',
    ]
    add_list(doc, arch)
    
    doc.add_page_break()
    
    add_title(doc, '第三章 用户角色与场景分析', level=1)
    
    add_title(doc, '3.1 用户角色', level=2)
    data = [
        ['角色', '描述', '使用场景'],
        ['系统管理员', '负责系统配置和维护', '配置数据源、推送渠道、定时任务'],
        ['AI 从业者', 'AI 研发工程师、产品经理', '获取 AI 技术动态和行业趋势'],
        ['企业团队', '科技公司技术团队', '团队内部共享技术资讯'],
        ['投资机构', '风险投资、私募股权从业者', '监控 AI 领域投资机会'],
        ['高校研究', '高校师生、科研机构人员', '跟踪最新研究进展'],
        ['个人用户', '科技爱好者、自媒体作者', '获取高质量的 AI 资讯'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.2 典型使用场景', level=2)
    
    scenarios = [
        {
            'id': 'S001',
            'name': '每日新闻简报推送',
            'description': '系统每天早上 8:00 自动爬取新闻，生成简报并推送到企业微信群',
            'actor': '企业团队',
            'flow': ['系统定时触发新闻爬取任务', '爬取多个数据源的新闻', '对新闻进行摘要和翻译', '生成新闻简报', '推送到企业微信群'],
        },
        {
            'id': 'S002',
            'name': '实时新闻浏览',
            'description': '用户通过 Web 前端界面查看最新新闻列表，支持筛选和搜索',
            'actor': 'AI 从业者',
            'flow': ['用户打开前端页面', '查看新闻列表', '筛选感兴趣的新闻', '点击查看新闻详情'],
        },
        {
            'id': 'S003',
            'name': '数据分析与监控',
            'description': '用户通过 Dashboard 查看新闻趋势、来源分布和关键词统计',
            'actor': '系统管理员',
            'flow': ['用户打开 Dashboard 页面', '查看新闻总数和今日新增', '查看近7天趋势图表', '查看来源分布和关键词统计'],
        },
        {
            'id': 'S004',
            'name': '手动触发爬取',
            'description': '系统管理员手动触发新闻爬取任务，获取最新新闻',
            'actor': '系统管理员',
            'flow': ['管理员登录系统', '点击手动爬取按钮', '系统执行爬取任务', '爬取完成后更新新闻列表'],
        },
        {
            'id': 'S005',
            'name': '配置推送渠道',
            'description': '系统管理员配置企业微信、钉钉、飞书的推送参数',
            'actor': '系统管理员',
            'flow': ['管理员进入设置页面', '填写推送渠道配置', '保存配置', '测试推送功能'],
        },
    ]
    
    for scenario in scenarios:
        add_paragraph(doc, f"**场景编号**：{scenario['id']}", bold=True)
        add_paragraph(doc, f"**场景名称**：{scenario['name']}")
        add_paragraph(doc, f"**场景描述**：{scenario['description']}")
        add_paragraph(doc, f"**参与者**：{scenario['actor']}")
        add_paragraph(doc, "**执行流程**：")
        for i, step in enumerate(scenario['flow'], 1):
            add_paragraph(doc, f"{i}. {step}", indent=True)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    add_title(doc, '第四章 功能需求', level=1)
    
    add_title(doc, '4.1 新闻爬虫模块', level=2)
    
    add_title(doc, '4.1.1 功能概述', level=3)
    add_paragraph(doc, '新闻爬虫模块负责从多个 RSS 数据源自动采集 AI 相关新闻，支持定时爬取和手动触发爬取。', indent=True)
    
    add_title(doc, '4.1.2 用例描述', level=3)
    
    add_usecase(doc, 'UC-001', '定时爬取新闻', '系统按照预定时间自动爬取新闻', '系统', '系统已启动，定时任务已配置', ['系统到达预定时间', '系统调用爬虫接口', '爬虫请求 RSS Feed', '解析 RSS 内容', '提取新闻标题、链接、内容、来源', '保存新闻数据'], '新闻数据已更新到数据库', ['网络异常：重试3次后记录错误日志', 'RSS Feed 格式错误：跳过该数据源'])
    
    doc.add_paragraph()
    
    add_usecase(doc, 'UC-002', '手动爬取新闻', '用户手动触发新闻爬取', '系统管理员', '用户已登录系统', ['用户点击手动爬取按钮', '系统调用爬虫接口', '爬虫请求 RSS Feed', '解析 RSS 内容', '提取新闻信息', '保存新闻数据', '返回爬取结果'], '新闻数据已更新，用户收到爬取结果通知', ['网络异常：显示错误提示', '爬虫服务未启动：提示用户启动服务'])
    
    doc.add_paragraph()
    
    add_title(doc, '4.1.3 功能需求列表', level=3)
    data = [
        ['需求编号', '需求描述', '优先级', '状态'],
        ['FR-001', '支持从多个 RSS 数据源采集新闻', '高', '已实现'],
        ['FR-002', '支持 TechCrunch AI 数据源', '高', '已实现'],
        ['FR-003', '支持 VentureBeat AI 数据源', '高', '已实现'],
        ['FR-004', '支持 MIT Technology Review 数据源', '中', '待实现'],
        ['FR-005', '支持 ZDNet AI 数据源', '中', '待实现'],
        ['FR-006', '支持 The Verge AI 数据源', '中', '待实现'],
        ['FR-007', '自动提取新闻标题、链接、内容、来源、日期', '高', '已实现'],
        ['FR-008', '支持自定义请求间隔，避免被封禁', '高', '已实现'],
        ['FR-009', '支持失败重试机制', '中', '已实现'],
        ['FR-010', '支持数据源配置管理', '中', '待实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.2 文本处理模块', level=2)
    
    add_title(doc, '4.2.1 功能概述', level=3)
    add_paragraph(doc, '文本处理模块负责对新闻内容进行摘要提取和翻译处理，将英文新闻转换为中文，方便用户阅读。', indent=True)
    
    add_title(doc, '4.2.2 用例描述', level=3)
    
    add_usecase(doc, 'UC-003', '提取新闻摘要', '系统对新闻内容进行摘要提取', '系统', '新闻内容已获取', ['系统接收新闻内容', '使用摘要算法提取关键信息', '生成200字以内的摘要', '保存摘要内容'], '摘要已保存到新闻数据中', ['新闻内容过短：直接使用原文作为摘要'])
    
    doc.add_paragraph()
    
    add_usecase(doc, 'UC-004', '翻译新闻内容', '系统将英文新闻翻译为中文', '系统', '新闻内容已获取', ['系统检测新闻语言', '若为英文，调用翻译服务', '翻译标题和内容', '保存翻译结果', '翻译失败时使用原文'], '翻译结果已保存到新闻数据中', ['翻译服务不可用：降级使用备用翻译服务'])
    
    doc.add_paragraph()
    
    add_title(doc, '4.2.3 功能需求列表', level=3)
    data = [
        ['需求编号', '需求描述', '优先级', '状态'],
        ['FR-011', '支持基于词频统计的摘要提取', '高', '已实现'],
        ['FR-012', '支持基于 TF-IDF 的摘要提取', '中', '待实现'],
        ['FR-013', '支持基于 transformer 的摘要提取', '低', '待实现'],
        ['FR-014', '支持中英文翻译', '高', '已实现'],
        ['FR-015', '支持百度翻译 API', '高', '待实现'],
        ['FR-016', '支持腾讯翻译 API', '中', '待实现'],
        ['FR-017', '支持 Google 翻译', '中', '已实现'],
        ['FR-018', '支持 MyMemory 翻译 API', '中', '已实现'],
        ['FR-019', '支持多层降级策略', '高', '已实现'],
        ['FR-020', '支持翻译缓存，避免重复翻译', '中', '待实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.3 推送模块', level=2)
    
    add_title(doc, '4.3.1 功能概述', level=3)
    add_paragraph(doc, '推送模块负责将新闻简报发送到企业微信、钉钉、飞书等平台，支持 Markdown 和纯文本格式。', indent=True)
    
    add_title(doc, '4.3.2 用例描述', level=3)
    
    add_usecase(doc, 'UC-005', '推送新闻简报', '系统将新闻简报推送到配置的渠道', '系统', '新闻简报已生成，推送配置已完成', ['系统获取推送配置', '生成推送内容', '调用企业微信接口推送', '调用钉钉接口推送', '调用飞书接口推送', '记录推送结果'], '推送结果已记录', ['推送失败：记录错误日志，不影响其他渠道'])
    
    doc.add_paragraph()
    
    add_title(doc, '4.3.3 功能需求列表', level=3)
    data = [
        ['需求编号', '需求描述', '优先级', '状态'],
        ['FR-021', '支持企业微信推送', '高', '已实现'],
        ['FR-022', '支持钉钉推送', '高', '已实现'],
        ['FR-023', '支持飞书推送', '高', '已实现'],
        ['FR-024', '支持 Markdown 格式推送', '高', '已实现'],
        ['FR-025', '支持纯文本格式推送', '中', '已实现'],
        ['FR-026', '支持推送配置管理', '高', '已实现'],
        ['FR-027', '支持推送测试', '中', '已实现'],
        ['FR-028', '支持推送结果记录', '中', '待实现'],
        ['FR-029', '支持推送失败重试', '中', '待实现'],
        ['FR-030', '支持自定义推送时间', '高', '已实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.4 定时任务模块', level=2)
    
    add_title(doc, '4.4.1 功能概述', level=3)
    add_paragraph(doc, '定时任务模块负责按预定时间自动执行新闻采集和推送任务，支持每天多个时间点运行。', indent=True)
    
    add_title(doc, '4.4.2 用例描述', level=3)
    
    add_usecase(doc, 'UC-006', '定时执行任务', '系统按照配置的时间自动执行任务', '系统', '定时任务已配置，系统已启动', ['系统检测当前时间', '匹配预定时间点', '执行新闻爬取任务', '执行新闻摘要任务', '执行新闻翻译任务', '执行新闻推送任务', '记录任务执行结果'], '任务执行完成，结果已记录', ['任务执行失败：记录错误日志，下次继续执行'])
    
    doc.add_paragraph()
    
    add_title(doc, '4.4.3 功能需求列表', level=3)
    data = [
        ['需求编号', '需求描述', '优先级', '状态'],
        ['FR-031', '支持每天多个时间点执行任务', '高', '已实现'],
        ['FR-032', '支持配置执行时间', '高', '已实现'],
        ['FR-033', '支持任务执行日志记录', '中', '已实现'],
        ['FR-034', '支持任务执行状态查询', '中', '待实现'],
        ['FR-035', '支持任务暂停和恢复', '低', '待实现'],
        ['FR-036', '支持任务执行失败告警', '中', '待实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.5 前端界面模块', level=2)
    
    add_title(doc, '4.5.1 功能概述', level=3)
    add_paragraph(doc, '前端界面模块提供用户交互界面，包括 Dashboard、新闻列表、数据分析和系统设置页面。', indent=True)
    
    add_title(doc, '4.5.2 功能需求列表', level=3)
    data = [
        ['需求编号', '需求描述', '优先级', '状态'],
        ['FR-037', 'Dashboard 页面：展示新闻总数、今日新增、趋势图表', '高', '已实现'],
        ['FR-038', '新闻列表页面：展示新闻列表，支持筛选和搜索', '高', '已实现'],
        ['FR-039', '新闻详情页面：展示新闻完整内容', '高', '已实现'],
        ['FR-040', '数据分析页面：来源分布、关键词统计', '高', '已实现'],
        ['FR-041', '系统设置页面：数据源配置、推送配置、定时任务配置', '高', '已实现'],
        ['FR-042', '支持响应式设计，适配不同屏幕尺寸', '中', '已实现'],
        ['FR-043', '支持新闻数据导出（Excel/CSV）', '中', '待实现'],
        ['FR-044', '支持新闻分类筛选', '中', '待实现'],
        ['FR-045', '支持新闻收藏功能', '低', '待实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.6 后端 API 模块', level=2)
    
    add_title(doc, '4.6.1 功能概述', level=3)
    add_paragraph(doc, '后端 API 模块提供 RESTful 接口，支持新闻管理、数据分析和系统配置。', indent=True)
    
    add_title(doc, '4.6.2 API 接口列表', level=3)
    data = [
        ['API 路径', 'HTTP 方法', '功能描述', '优先级', '状态'],
        ['/api/news', 'GET', '获取新闻列表', '高', '已实现'],
        ['/api/news', 'POST', '爬取并保存新闻', '高', '已实现'],
        ['/api/news/<id>', 'GET', '获取新闻详情', '高', '已实现'],
        ['/api/news/<id>', 'DELETE', '删除新闻', '中', '已实现'],
        ['/api/analytics/stats', 'GET', '获取统计数据', '高', '已实现'],
        ['/api/analytics/trend', 'GET', '获取趋势数据', '高', '已实现'],
        ['/api/config', 'GET', '获取系统配置', '高', '已实现'],
        ['/api/config', 'POST', '更新系统配置', '高', '已实现'],
        ['/api/push', 'POST', '手动触发推送', '中', '已实现'],
        ['/api/push/test', 'POST', '测试推送配置', '中', '已实现'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第五章 非功能需求', level=1)
    
    add_title(doc, '5.1 性能需求', level=2)
    data = [
        ['需求编号', '需求描述', '指标'],
        ['NFR-001', '新闻爬取速度', '单个数据源爬取时间 ≤ 10 秒'],
        ['NFR-002', '翻译响应时间', '单篇新闻翻译时间 ≤ 5 秒'],
        ['NFR-003', 'API 响应时间', 'API 请求响应时间 ≤ 1 秒'],
        ['NFR-004', '前端页面加载时间', '页面首次加载时间 ≤ 3 秒'],
        ['NFR-005', '系统吞吐量', '支持同时处理 100+ 条新闻'],
        ['NFR-006', '并发处理能力', '支持 10+ 并发用户访问'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.2 安全性需求', level=2)
    data = [
        ['需求编号', '需求描述', '说明'],
        ['NFR-007', '推送配置加密存储', 'Webhook URL 和 Token 加密存储'],
        ['NFR-008', 'API 访问控制', '支持 IP 白名单或 API Key 认证'],
        ['NFR-009', '输入数据校验', '对用户输入进行严格校验，防止注入攻击'],
        ['NFR-010', '日志安全', '日志中不记录敏感信息'],
        ['NFR-011', 'HTTPS 支持', '生产环境支持 HTTPS 加密传输'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.3 可用性需求', level=2)
    data = [
        ['需求编号', '需求描述', '指标'],
        ['NFR-012', '系统可用性', '系统可用率 ≥ 99%'],
        ['NFR-013', '故障恢复时间', '系统故障恢复时间 ≤ 15 分钟'],
        ['NFR-014', '数据备份', '每日自动备份数据'],
        ['NFR-015', '错误处理', '系统错误率 ≤ 1%'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.4 兼容性需求', level=2)
    data = [
        ['需求编号', '需求描述', '说明'],
        ['NFR-016', '浏览器兼容性', '支持 Chrome 80+、Firefox 75+、Edge 80+'],
        ['NFR-017', '操作系统兼容性', '支持 Windows 10/11、Linux、macOS'],
        ['NFR-018', '移动端兼容性', '支持手机浏览器访问（响应式设计）'],
        ['NFR-019', 'Python 版本兼容性', '支持 Python 3.10+'],
        ['NFR-020', 'Node.js 版本兼容性', '支持 Node.js 18+'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.5 可维护性需求', level=2)
    data = [
        ['需求编号', '需求描述', '说明'],
        ['NFR-021', '代码可读性', '代码符合 PEP 8 规范，有完善的注释'],
        ['NFR-022', '模块化设计', '系统按功能模块划分，模块间低耦合'],
        ['NFR-023', '日志系统', '完善的日志记录，便于问题排查'],
        ['NFR-024', '监控告警', '支持关键指标监控和异常告警'],
        ['NFR-025', '文档完善', '提供详细的开发文档和用户手册'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第六章 数据需求', level=1)
    
    add_title(doc, '6.1 数据实体', level=2)
    
    add_title(doc, '6.1.1 新闻数据', level=3)
    data = [
        ['字段名', '类型', '长度', '必填', '描述'],
        ['id', 'string', '50', '是', '新闻唯一标识'],
        ['title', 'string', '500', '是', '新闻标题（英文）'],
        ['title_zh', 'string', '500', '否', '新闻标题（中文）'],
        ['link', 'string', '1000', '是', '新闻链接'],
        ['summary', 'string', '500', '否', '新闻摘要（英文）'],
        ['summary_zh', 'string', '500', '否', '新闻摘要（中文）'],
        ['content', 'string', '-', '是', '新闻内容'],
        ['source', 'string', '100', '是', '新闻来源'],
        ['date', 'string', '20', '是', '新闻发布日期'],
        ['tags', 'array', '-', '否', '新闻标签'],
        ['created_at', 'string', '20', '是', '记录创建时间'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '6.1.2 系统配置', level=3)
    data = [
        ['字段名', '类型', '长度', '必填', '描述'],
        ['wechat_webhook', 'string', '500', '否', '企业微信 Webhook URL'],
        ['dingtalk_token', 'string', '100', '否', '钉钉机器人 Token'],
        ['feishu_webhook', 'string', '500', '否', '飞书 Webhook URL'],
        ['article_count', 'integer', '-', '是', '推送文章数量'],
        ['schedule_times', 'array', '-', '是', '定时任务时间'],
        ['rss_feeds', 'array', '-', '是', 'RSS 数据源配置'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '6.1.3 定时任务配置', level=3)
    data = [
        ['字段名', '类型', '长度', '必填', '描述'],
        ['enabled', 'boolean', '-', '是', '是否启用'],
        ['times', 'array', '-', '是', '执行时间列表'],
        ['last_run', 'string', '20', '否', '上次执行时间'],
        ['next_run', 'string', '20', '否', '下次执行时间'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '6.2 数据存储', level=2)
    data = [
        ['数据类型', '存储方式', '存储位置', '说明'],
        ['新闻数据', 'JSON 文件', 'data/articles.json', '存储新闻列表'],
        ['系统配置', 'JSON 文件', 'data/config.json', '存储系统配置'],
        ['日志数据', '文本文件', 'logs/', '存储系统日志'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '6.3 数据备份', level=2)
    data = [
        ['备份项', '备份频率', '备份方式', '保留期限'],
        ['新闻数据', '每日', '自动备份', '30天'],
        ['系统配置', '每次修改', '自动备份', '永久'],
        ['日志数据', '每日', '自动备份', '7天'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第七章 接口需求', level=1)
    
    add_title(doc, '7.1 内部接口', level=2)
    
    add_title(doc, '7.1.1 爬虫接口', level=3)
    data = [
        ['接口名称', '输入参数', '输出参数', '说明'],
        ['crawl_ai_news()', '无', 'articles: list', '爬取所有配置的数据源'],
        ['fetch_rss_feed(url)', 'url: string', 'feed: dict', '获取单个 RSS Feed'],
        ['parse_rss_content(content)', 'content: string', 'items: list', '解析 RSS 内容'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.1.2 文本处理接口', level=3)
    data = [
        ['接口名称', '输入参数', '输出参数', '说明'],
        ['extract_summary(text)', 'text: string', 'summary: string', '提取文本摘要'],
        ['translate_text(text, src, dst)', 'text, src, dst: string', 'result: string', '翻译文本'],
        ['extract_tags(text)', 'text: string', 'tags: list', '提取关键词标签'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.1.3 推送接口', level=3)
    data = [
        ['接口名称', '输入参数', '输出参数', '说明'],
        ['send_wechat(content, url)', 'content, url: string', 'success: boolean', '发送企业微信消息'],
        ['send_dingtalk(content, token)', 'content, token: string', 'success: boolean', '发送钉钉消息'],
        ['send_feishu(content, url)', 'content, url: string', 'success: boolean', '发送飞书消息'],
        ['send_all(content)', 'content: string', 'results: dict', '发送所有配置的渠道'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.2 外部接口', level=2)
    
    add_title(doc, '7.2.1 第三方 API', level=3)
    data = [
        ['API 名称', '接口地址', '用途', '认证方式'],
        ['企业微信机器人', 'https://qyapi.weixin.qq.com/cgi-bin/webhook/send', '消息推送', 'Webhook Key'],
        ['钉钉机器人', 'https://oapi.dingtalk.com/robot/send', '消息推送', 'Token + 签名'],
        ['飞书机器人', 'https://open.feishu.cn/open-apis/bot/v2/hook/{hook_id}', '消息推送', 'Webhook URL'],
        ['百度翻译 API', 'https://fanyi-api.baidu.com/api/trans/vip/translate', '文本翻译', 'API Key'],
        ['腾讯翻译 API', 'https://tmt.tencentcloudapi.com/', '文本翻译', 'Secret ID/Key'],
        ['Google 翻译', 'https://translate.google.com/', '文本翻译', '无需认证'],
        ['MyMemory API', 'https://api.mymemory.translated.net/get', '文本翻译', '无需认证'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.2.2 RSS 数据源', level=3)
    data = [
        ['数据源名称', 'RSS 地址', '内容类型'],
        ['TechCrunch AI', 'https://techcrunch.com/tag/ai/feed/', 'AI 新闻'],
        ['VentureBeat AI', 'https://venturebeat.com/category/ai/feed/', 'AI 新闻'],
        ['MIT Technology Review', 'https://www.technologyreview.com/feed/', '科技新闻'],
        ['ZDNet AI', 'https://www.zdnet.com/topic/artificial-intelligence/rss.xml', 'AI 新闻'],
        ['The Verge', 'https://www.theverge.com/rss/index.xml', '科技新闻'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.3 前后端接口', level=2)
    
    add_title(doc, '7.3.1 新闻管理接口', level=3)
    data = [
        ['API 路径', '方法', '请求体', '响应体', '说明'],
        ['/api/news', 'GET', 'page, limit', '{status, data: articles}', '分页获取新闻'],
        ['/api/news', 'POST', '无', '{status, message, count}', '爬取并保存新闻'],
        ['/api/news/:id', 'GET', '无', '{status, data: article}', '获取新闻详情'],
        ['/api/news/:id', 'DELETE', '无', '{status, message}', '删除新闻'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.3.2 数据分析接口', level=3)
    data = [
        ['API 路径', '方法', '请求体', '响应体', '说明'],
        ['/api/analytics/stats', 'GET', '无', '{status, data: stats}', '获取统计数据'],
        ['/api/analytics/trend', 'GET', 'days', '{status, data: trend}', '获取趋势数据'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '7.3.3 配置管理接口', level=3)
    data = [
        ['API 路径', '方法', '请求体', '响应体', '说明'],
        ['/api/config', 'GET', '无', '{status, data: config}', '获取系统配置'],
        ['/api/config', 'POST', 'config', '{status, message}', '更新系统配置'],
        ['/api/push', 'POST', 'content', '{status, message}', '手动触发推送'],
        ['/api/push/test', 'POST', '无', '{status, message}', '测试推送配置'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第八章 约束条件', level=1)
    
    add_title(doc, '8.1 技术约束', level=2)
    constraints = [
        '后端语言：Python 3.10+',
        '前端框架：React 18 + TypeScript',
        'Web 框架：Flask 3',
        '构建工具：Vite 6',
        '样式框架：TailwindCSS 3',
        '图表库：Chart.js',
        '打包工具：PyInstaller',
        '数据库：JSON 文件 / SQLite',
        '操作系统：Windows 10/11、Linux、macOS',
    ]
    add_list(doc, constraints)
    
    add_title(doc, '8.2 资源约束', level=2)
    constraints = [
        '服务器最低配置：CPU 2核，内存 4GB，存储 50GB',
        '网络带宽：至少 1Mbps',
        '翻译服务免费额度：每日有限制',
        '爬虫请求频率：遵守目标网站规则，避免被封禁',
    ]
    add_list(doc, constraints)
    
    add_title(doc, '8.3 时间约束', level=2)
    constraints = [
        '新闻爬取时间：建议在凌晨或低峰期执行，避免影响系统性能',
        '推送时间：建议在工作时间推送，确保用户能及时查看',
        '数据更新频率：每天至少更新一次',
    ]
    add_list(doc, constraints)
    
    add_title(doc, '8.4 法律约束', level=2)
    constraints = [
        '数据采集需遵守目标网站的 robots.txt 协议',
        '新闻内容仅限个人/内部参考使用，不得用于商业分发',
        '用户配置信息需本地存储，不得传输到第三方',
        '遵守相关隐私保护法律法规',
    ]
    add_list(doc, constraints)
    
    doc.add_page_break()
    
    add_title(doc, '第九章 验收标准', level=1)
    
    add_title(doc, '9.1 功能验收标准', level=2)
    data = [
        ['需求编号', '验收标准'],
        ['FR-001', '系统能从至少2个 RSS 数据源成功采集新闻'],
        ['FR-007', '新闻数据包含标题、链接、内容、来源、日期字段'],
        ['FR-011', '摘要长度不超过200字，能准确概括新闻核心内容'],
        ['FR-014', '英文新闻能成功翻译为中文'],
        ['FR-021', '能成功将新闻简报推送到企业微信'],
        ['FR-022', '能成功将新闻简报推送到钉钉'],
        ['FR-023', '能成功将新闻简报推送到飞书'],
        ['FR-031', '定时任务能在预定时间自动执行'],
        ['FR-037', 'Dashboard 能展示新闻总数、今日新增和趋势图表'],
        ['FR-038', '新闻列表能正确展示，支持筛选和搜索'],
        ['FR-041', '系统设置页面能配置数据源、推送和定时任务'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '9.2 非功能验收标准', level=2)
    data = [
        ['需求编号', '验收标准'],
        ['NFR-003', 'API 请求响应时间 ≤ 1 秒'],
        ['NFR-004', '前端页面首次加载时间 ≤ 3 秒'],
        ['NFR-007', '推送配置信息加密存储'],
        ['NFR-012', '系统可用率 ≥ 99%'],
        ['NFR-016', '在 Chrome 80+、Firefox 75+、Edge 80+ 中正常运行'],
        ['NFR-017', '在 Windows 10/11、Linux、macOS 中正常运行'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '9.3 接口验收标准', level=2)
    data = [
        ['接口名称', '验收标准'],
        ['/api/news GET', '返回新闻列表，支持分页'],
        ['/api/news POST', '成功爬取新闻并返回爬取数量'],
        ['/api/analytics/stats', '返回统计数据，包含趋势、来源分布、关键词'],
        ['/api/config GET', '返回系统配置信息'],
        ['/api/config POST', '成功更新系统配置'],
        ['企业微信推送', '消息能成功发送到企业微信群'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('— 文档结束 —')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = os.path.join(os.path.dirname(__file__), 'AI新闻爬虫系统需求规格说明书.docx')
    doc.save(output_path)
    print(f"需求规格说明书已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_requirements_doc()
