from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
import os
import datetime

def add_title(doc, text, level=1):
    if level == 1:
        style = 'Heading 1'
    elif level == 2:
        style = 'Heading 2'
    elif level == 3:
        style = 'Heading 3'
    else:
        style = 'Heading 4'
    
    title = doc.add_heading(text, level=level)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = '微软雅黑'
    title.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return title

def add_paragraph(doc, text, font_size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
    return p

def add_table(doc, data, headers=None):
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if headers and i == 0:
                cell.text = headers[j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                cell.text = data[i - 1 if headers else i][j]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    table.style = 'Table Grid'
    return table

def add_list(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        if ordered:
            p = doc.add_paragraph(f"{i}. {item}", style='List Number')
        else:
            p = doc.add_paragraph(f"• {item}", style='List Bullet')
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def generate_report():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    add_title(doc, 'AI新闻爬虫系统', level=1)
    add_title(doc, '可行性分析与项目开发计划报告', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run(f'报告生成时间：{datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run('报告版本：v1.0')
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    add_title(doc, '目录', level=1)
    toc_items = [
        ('第一章 项目概述', '1'),
        ('第二章 可行性分析', '5'),
        ('第三章 项目开发计划', '20'),
        ('第四章 风险评估与应对', '35'),
        ('第五章 投资估算与效益分析', '42'),
        ('第六章 结论与建议', '48'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}'.ljust(40) + f'第 {page} 页')
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    add_title(doc, '第一章 项目概述', level=1)
    
    add_title(doc, '1.1 项目背景', level=2)
    add_paragraph(doc, '随着人工智能技术的快速发展，AI 相关新闻资讯呈爆炸式增长。企业、团队和个人需要及时了解 AI 领域的最新动态，但人工浏览多个科技媒体网站效率低下且容易遗漏重要信息。本项目旨在开发一套自动化的 AI 新闻采集、分析与推送系统，帮助用户高效获取 AI 领域的最新资讯。', indent=True)
    
    add_title(doc, '1.2 项目目标', level=2)
    add_paragraph(doc, '本项目的核心目标是构建一个完整的 AI 新闻自动化处理平台，具体包括：', indent=True)
    goals = [
        '自动化采集：从国际主流科技媒体自动抓取 AI 相关新闻',
        '智能分析：对新闻内容进行摘要提取和翻译处理',
        '多端推送：通过企业微信、钉钉、飞书等平台推送新闻简报',
        '实时监控：提供 Web 前端界面进行实时动态分析',
        '定时运行：支持定时任务自动执行新闻采集和推送',
    ]
    add_list(doc, goals, ordered=True)
    
    add_title(doc, '1.3 项目现状', level=2)
    add_paragraph(doc, '目前项目已完成第一阶段开发，具备以下功能：', indent=True)
    status = [
        '✅ 新闻爬虫模块：支持 TechCrunch AI、VentureBeat AI 两个 RSS 数据源',
        '✅ 文本摘要模块：基于词频统计的摘要提取算法',
        '✅ 翻译模块：支持百度、Google、腾讯翻译及 MyMemory API',
        '✅ 推送模块：支持企业微信、钉钉、飞书多平台推送',
        '✅ 定时任务：基于标准库的定时任务调度器',
        '✅ Flask 后端 API：提供新闻管理、分析和配置接口',
        '✅ React 前端界面：Dashboard、新闻列表、数据分析、系统设置页面',
        '✅ 数据可视化：趋势图表、来源分布、关键词统计',
    ]
    add_list(doc, status)
    
    add_title(doc, '1.4 目标用户', level=2)
    data = [
        ['用户群体', '描述', '需求痛点'],
        ['AI 从业者', 'AI 研发工程师、产品经理、研究员', '需要及时了解技术动态和行业趋势'],
        ['企业团队', '科技公司技术团队、产品团队', '需要团队内部共享技术资讯'],
        ['投资机构', '风险投资、私募股权从业者', '需要监控 AI 领域投资机会'],
        ['高校研究', '高校师生、科研机构人员', '需要跟踪最新研究进展'],
        ['个人用户', '科技爱好者、自媒体作者', '需要获取高质量的 AI 资讯'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第二章 可行性分析', level=1)
    
    add_title(doc, '2.1 技术可行性', level=2)
    add_paragraph(doc, '本项目所采用的技术方案成熟可行，所有技术栈均为业界主流选择，具备良好的技术支撑和社区生态。', indent=True)
    
    add_title(doc, '2.1.1 技术栈评估', level=3)
    data = [
        ['技术领域', '技术选型', '成熟度', '可行性评估'],
        ['后端语言', 'Python 3.13', '高', '成熟稳定，生态完善'],
        ['Web框架', 'Flask 3', '高', '轻量级，适合快速开发'],
        ['前端框架', 'React 18 + TypeScript', '高', '组件化开发，类型安全'],
        ['构建工具', 'Vite 6', '高', '快速热更新，开发体验好'],
        ['样式框架', 'TailwindCSS 3', '高', '原子化CSS，快速开发'],
        ['图表库', 'Chart.js', '高', '成熟稳定，文档完善'],
        ['爬虫技术', 'requests + BeautifulSoup', '高', 'Python生态标准方案'],
        ['翻译服务', 'translators库', '中', '免费服务，稳定性有限'],
        ['打包工具', 'PyInstaller', '高', '成熟稳定，支持单文件打包'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.1.2 技术难点分析', level=3)
    data = [
        ['技术难点', '风险等级', '解决方案', '可行性'],
        ['反爬机制规避', '中', '设置 User-Agent、随机延迟、请求频率控制', '可行'],
        ['翻译质量保障', '中', '多层降级策略，优先使用高质量翻译服务', '可行'],
        ['数据持久化', '低', '使用 JSON 文件或 SQLite 数据库', '可行'],
        ['定时任务可靠性', '低', '基于标准库实现，支持异常重试', '可行'],
        ['前端性能优化', '中', '组件懒加载、数据缓存、代码分割', '可行'],
        ['跨域问题', '低', 'Flask-CORS 中间件', '可行'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.1.3 技术可行性结论', level=3)
    add_paragraph(doc, '综合评估，本项目的技术方案完全可行。所有核心技术均为成熟方案，技术难点均有明确的解决方案。项目采用的前后端分离架构、模块化设计等技术选型符合现代 Web 开发最佳实践，能够支撑系统的功能需求和扩展需求。', indent=True)
    
    add_title(doc, '2.2 经济可行性', level=2)
    add_paragraph(doc, '本项目的经济投入主要包括开发成本、运维成本和潜在收益，以下从成本和收益两个维度进行分析。', indent=True)
    
    add_title(doc, '2.2.1 成本分析', level=3)
    data = [
        ['成本类型', '项目', '预估金额', '说明'],
        ['开发成本', '人力成本（开发）', '3-5万元', '1名全栈开发，约1-2个月'],
        ['开发成本', '人力成本（测试）', '0.5-1万元', '测试和调试'],
        ['开发成本', '工具费用', '0', '使用开源工具'],
        ['运维成本', '服务器费用', '0-500元/月', '可使用本地服务器或云服务器'],
        ['运维成本', '翻译服务费用', '0-100元/月', '免费额度足够初期使用'],
        ['运维成本', '域名和SSL', '0-100元/年', '可选'],
        ['总计', '-', '约3-6万元', '初期投入'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.2.2 收益分析', level=3)
    data = [
        ['收益类型', '具体内容', '预估价值', '实现周期'],
        ['效率提升', '减少人工浏览时间，每人每天节省约30分钟', '每人每月约600元', '立即'],
        ['决策支持', '及时获取行业动态，辅助商业决策', '难以量化', '短期'],
        ['团队协作', '团队内部共享资讯，促进知识交流', '提高团队效率10-20%', '短期'],
        ['内容创作', '为自媒体、公众号提供内容素材', '潜在商业价值', '中期'],
        ['产品增值', '可作为企业内部工具，提升员工体验', '提升企业形象', '中期'],
        ['数据资产', '积累的新闻数据可用于分析和挖掘', '数据价值随时间增长', '长期'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.2.3 投资回报率', level=3)
    add_paragraph(doc, '根据成本收益分析，本项目具有良好的投资回报率：', indent=True)
    roi = [
        '短期（1-3个月）：覆盖开发成本，开始产生效率提升价值',
        '中期（3-6个月）：团队效率提升明显，开始产生商业价值',
        '长期（6个月以上）：数据资产增值，可拓展商业化服务',
        '静态投资回收期：约3-6个月',
        '动态投资回收期：约6-12个月',
    ]
    add_list(doc, roi)
    
    add_title(doc, '2.2.4 经济可行性结论', level=3)
    add_paragraph(doc, '本项目初期投入较低（约3-6万元），主要成本为人力成本，运维成本极低。项目带来的效率提升和决策支持价值显著，投资回收期较短。对于企业用户，团队效率提升带来的价值远超过开发成本；对于个人用户，可免费使用，无经济负担。因此，本项目在经济上完全可行。', indent=True)
    
    add_title(doc, '2.3 运营可行性', level=2)
    add_paragraph(doc, '运营可行性主要评估项目上线后的运维难度、用户接受度和运营成本。', indent=True)
    
    add_title(doc, '2.3.1 运维难度评估', level=3)
    data = [
        ['运维项目', '难度等级', '说明', '解决方案'],
        ['服务器部署', '低', '可使用 Windows 本地服务器或云服务器', '提供部署脚本和文档'],
        ['服务监控', '低', '需要监控服务运行状态', '使用 Windows 任务计划或云监控'],
        ['数据备份', '低', '定期备份数据文件', '自动化备份脚本'],
        ['版本更新', '中', '需要更新程序版本', '一键更新机制'],
        ['故障排查', '中', '需要处理网络异常、API 故障等', '完善的日志系统'],
        ['配置管理', '低', '管理推送配置、定时任务等', 'Web 前端配置界面'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.3.2 用户接受度评估', level=3)
    data = [
        ['用户群体', '接受度', '原因', '推广策略'],
        ['AI 从业者', '高', '专业需求强，主动获取资讯', '技术社区推广'],
        ['企业团队', '高', '团队协作需求', '企业内部推广'],
        ['投资机构', '中', '需要定制化功能', '提供定制服务'],
        ['高校研究', '中', '学术需求为主', '学术社区推广'],
        ['个人用户', '高', '免费使用，门槛低', '社交媒体推广'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.3.3 运营可行性结论', level=3)
    add_paragraph(doc, '本项目运维难度较低，主要运维工作包括服务器管理、数据备份和故障排查，均可通过自动化脚本和工具实现。用户接受度高，尤其是 AI 从业者和企业团队。运营成本低，无需专职运维人员，适合小型团队或个人运营。因此，本项目在运营上完全可行。', indent=True)
    
    add_title(doc, '2.4 法律可行性', level=2)
    add_paragraph(doc, '法律可行性主要评估项目在数据采集、内容使用、用户隐私等方面的法律合规性。', indent=True)
    
    add_title(doc, '2.4.1 数据采集合规性', level=3)
    data = [
        ['数据来源', '合规性', '说明', '风险等级'],
        ['RSS Feed', '合规', '公开的 RSS 订阅源，符合网站协议', '低'],
        ['新闻内容', '合规', '仅用于个人/内部参考，不用于商业分发', '低'],
        ['用户配置', '合规', '用户自行配置推送渠道', '低'],
        ['推送内容', '合规', '推送内容来源于公开媒体', '低'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.4.2 法律风险与应对', level=3)
    data = [
        ['风险类型', '风险描述', '应对措施', '风险等级'],
        ['版权风险', '新闻内容版权归属原媒体', '仅用于个人/内部参考，注明来源', '低'],
        ['隐私风险', '用户配置信息存储', '本地存储，不传输第三方', '低'],
        ['合规风险', '爬虫行为可能违反网站规则', '遵守 robots.txt，控制请求频率', '低'],
        ['数据安全', '配置文件包含敏感信息', '本地加密存储', '低'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '2.4.3 法律可行性结论', level=3)
    add_paragraph(doc, '本项目的数据采集来源为公开的 RSS Feed，符合网站协议。新闻内容仅用于个人或企业内部参考，不进行商业分发，符合版权法规定。用户配置信息本地存储，不传输到第三方，保护用户隐私。因此，本项目在法律上完全可行。', indent=True)
    
    add_title(doc, '2.5 可行性综合评估', level=2)
    data = [
        ['评估维度', '可行性', '权重', '得分'],
        ['技术可行性', '完全可行', '30%', '95分'],
        ['经济可行性', '完全可行', '25%', '90分'],
        ['运营可行性', '完全可行', '25%', '85分'],
        ['法律可行性', '完全可行', '20%', '95分'],
        ['综合评分', '-', '100%', '91分'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_paragraph(doc, '综合评估结果：本项目在技术、经济、运营和法律四个维度均完全可行，综合得分为 91 分（百分制），属于「高度可行」级别。建议立即启动后续开发工作。', indent=True)
    
    doc.add_page_break()
    
    add_title(doc, '第三章 项目开发计划', level=1)
    
    add_title(doc, '3.1 项目阶段划分', level=2)
    add_paragraph(doc, '本项目分为三个开发阶段，每个阶段有明确的目标和交付物：', indent=True)
    
    data = [
        ['阶段', '名称', '时间', '目标', '交付物'],
        ['第一阶段', '基础功能实现', '已完成', '实现核心爬虫、摘要、翻译、推送功能', '爬虫模块、推送模块、定时任务'],
        ['第二阶段', '前端界面开发', '已完成', '开发 Web 前端界面，实现实时动态分析', 'Dashboard、新闻列表、数据分析、设置页面'],
        ['第三阶段', '功能优化与扩展', '2-3周', '优化现有功能，添加新特性', '性能优化、数据源扩展、智能分类'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.2 第三阶段开发计划', level=2)
    add_paragraph(doc, '基于现有基础，第三阶段主要进行功能优化和扩展，计划周期为 2-3 周。', indent=True)
    
    add_title(doc, '3.2.1 功能优化', level=3)
    data = [
        ['序号', '优化项', '描述', '优先级', '预估工时'],
        ['1', '翻译质量提升', '接入商用翻译 API（百度/腾讯），提高翻译准确率', '高', '2天'],
        ['2', '摘要算法改进', '引入 TF-IDF 或基于 transformer 的摘要模型', '高', '3天'],
        ['3', '数据去重', '基于链接和标题的智能去重机制', '高', '2天'],
        ['4', '性能优化', '异步爬取、缓存机制、批量处理', '中', '3天'],
        ['5', '日志系统', '完善的文件日志和日志级别控制', '中', '2天'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.2.2 功能扩展', level=3)
    data = [
        ['序号', '扩展项', '描述', '优先级', '预估工时'],
        ['1', '数据源扩展', '添加 MIT Technology Review、ZDNet、The Verge 等数据源', '高', '3天'],
        ['2', '智能分类', '基于关键词的新闻自动分类（技术、商业、伦理等）', '中', '3天'],
        ['3', '数据导出', '支持导出为 Excel、CSV 格式', '中', '2天'],
        ['4', '多语言支持', '支持日语、韩语等其他语言翻译', '低', '2天'],
        ['5', '个性化推荐', '基于用户阅读历史的智能推荐', '低', '3天'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.2.3 开发时间线', level=3)
    timeline = """第1周：
  ├── 第1-2天：翻译质量提升
  ├── 第3-4天：数据去重功能
  └── 第5天：日志系统完善

第2周：
  ├── 第1-3天：摘要算法改进
  ├── 第4-5天：数据源扩展（新增3个数据源）

第3周：
  ├── 第1-2天：性能优化
  ├── 第3-4天：智能分类功能
  └── 第5天：测试与文档完善"""
    doc.add_paragraph()
    run = doc.paragraphs[-1].add_run(timeline)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    
    add_title(doc, '3.3 资源需求', level=2)
    data = [
        ['资源类型', '需求', '说明'],
        ['人力', '1名全栈开发工程师', '负责后端和前端开发'],
        ['硬件', '1台服务器（最低配置）', 'CPU: 2核, 内存: 4GB, 存储: 50GB'],
        ['软件', 'Windows/Linux 操作系统', 'Python 3.10+, Node.js 18+'],
        ['网络', '稳定的网络连接', '用于爬取新闻和推送消息'],
        ['工具', 'Git、VS Code、PyCharm', '开发工具'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.4 里程碑计划', level=2)
    data = [
        ['里程碑', '时间', '目标', '验收标准'],
        ['M1', '已完成', '基础功能上线', '爬虫、摘要、翻译、推送功能正常'],
        ['M2', '已完成', '前端界面上线', 'Dashboard、新闻列表、数据分析、设置页面正常'],
        ['M3', '第2周结束', '核心优化完成', '翻译准确率提升、数据去重、日志系统'],
        ['M4', '第3周结束', '功能扩展完成', '新增数据源、智能分类、数据导出'],
        ['M5', '第3周结束', '测试验收', '所有功能测试通过，文档完善'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '3.5 开发流程', level=2)
    add_paragraph(doc, '项目采用敏捷开发流程，具体如下：', indent=True)
    flow = [
        '需求分析 → 技术设计 → 编码实现 → 单元测试 → 集成测试 → 部署上线',
        '每周进行一次迭代回顾，评估进度和质量',
        '使用 Git 进行版本控制，遵循 Git Flow 工作流',
        '代码提交前进行代码审查，确保代码质量',
        '使用 pytest 进行后端测试，使用 React Testing Library 进行前端测试',
    ]
    add_list(doc, flow)
    
    doc.add_page_break()
    
    add_title(doc, '第四章 风险评估与应对', level=1)
    
    add_title(doc, '4.1 技术风险', level=2)
    data = [
        ['风险描述', '发生概率', '影响程度', '应对措施'],
        ['爬虫被封禁', '中', '高', '设置合理请求间隔，使用代理IP池，遵守网站规则'],
        ['翻译服务不可用', '中', '中', '多层降级策略，备用翻译服务，本地缓存'],
        ['API 接口变更', '低', '高', '定期检查 API 状态，预留兼容层'],
        ['数据格式变化', '中', '中', '增加数据校验，灵活的解析逻辑'],
        ['性能瓶颈', '中', '中', '异步处理，缓存机制，数据库优化'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.2 运营风险', level=2)
    data = [
        ['风险描述', '发生概率', '影响程度', '应对措施'],
        ['服务器故障', '低', '高', '定期备份数据，使用云服务器自动恢复'],
        ['网络中断', '中', '中', '本地缓存机制，网络恢复后自动同步'],
        ['用户流失', '中', '中', '持续优化功能，提供优质体验，收集用户反馈'],
        ['运营成本增加', '低', '低', '优化资源使用，使用低成本方案'],
        ['竞品出现', '中', '中', '持续创新，差异化竞争，快速迭代'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.3 市场风险', level=2)
    data = [
        ['风险描述', '发生概率', '影响程度', '应对措施'],
        ['需求变化', '中', '中', '保持灵活性，快速响应市场变化'],
        ['技术替代', '低', '高', '关注技术趋势，持续学习和创新'],
        ['政策法规变化', '低', '高', '关注相关政策，确保合规运营'],
        ['经济环境变化', '低', '中', '控制成本，保持现金流'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '4.4 风险应对策略', level=2)
    strategies = [
        '风险预防：在开发阶段充分考虑各种风险，设计容错机制',
        '风险转移：对于高风险项，考虑使用成熟的第三方服务',
        '风险缓解：制定应急预案，定期进行风险评估和演练',
        '风险监控：建立监控系统，及时发现和处理风险',
        '风险恢复：建立备份和恢复机制，确保业务连续性',
    ]
    add_list(doc, strategies)
    
    doc.add_page_break()
    
    add_title(doc, '第五章 投资估算与效益分析', level=1)
    
    add_title(doc, '5.1 投资估算', level=2)
    data = [
        ['投资阶段', '项目', '金额', '说明'],
        ['初期投资', '开发人力', '3-5万元', '1名全栈开发，1-2个月'],
        ['初期投资', '测试人力', '0.5-1万元', '测试和调试'],
        ['初期投资', '服务器费用', '0-0.6万元/年', '云服务器或本地服务器'],
        ['初期投资', '其他费用', '0.5万元', '域名、SSL证书等'],
        ['总计', '-', '约4-7万元', '初期总投资'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.2 运营成本', level=2)
    data = [
        ['成本项目', '年费用', '说明'],
        ['服务器费用', '0-6000元', '视服务器配置而定'],
        ['翻译服务费用', '0-1200元', '免费额度足够，超出部分付费'],
        ['域名和SSL', '0-1000元', '可选'],
        ['维护人力', '2-4万元', '兼职维护，约200小时/年'],
        ['总计', '约2-5万元/年', '年运营成本'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.3 效益分析', level=2)
    data = [
        ['效益类型', '量化指标', '年度价值'],
        ['效率提升', '节省人工浏览时间，每人每天30分钟', '每人每年约1.2万元'],
        ['决策支持', '及时获取行业动态，辅助商业决策', '难以量化，价值显著'],
        ['团队协作', '团队效率提升10-20%', '团队每年约5-10万元'],
        ['内容创作', '提供内容素材', '视商业化程度而定'],
        ['数据资产', '积累新闻数据', '价值随时间增长'],
    ]
    add_table(doc, data[1:], data[0])
    
    add_title(doc, '5.4 投资回报率分析', level=2)
    add_paragraph(doc, '基于以上成本和效益分析，本项目的投资回报率计算如下：', indent=True)
    
    roi_data = [
        ['指标', '数值'],
        ['初期投资', '约4-7万元'],
        ['年运营成本', '约2-5万元'],
        ['年效益（按10人团队计算）', '约12-22万元'],
        ['年净利润', '约5-15万元'],
        ['投资回报率', '约70-200%'],
        ['静态回收期', '约3-6个月'],
        ['动态回收期', '约6-12个月'],
    ]
    add_table(doc, roi_data[1:], roi_data[0])
    
    add_title(doc, '5.5 敏感性分析', level=2)
    data = [
        ['变量', '基准值', '变动幅度', '对ROI的影响'],
        ['开发成本', '5万元', '+20%', 'ROI下降约15%'],
        ['开发成本', '5万元', '-20%', 'ROI上升约15%'],
        ['团队规模', '10人', '+50%', '效益增加约50%'],
        ['团队规模', '10人', '-50%', '效益减少约50%'],
        ['运营成本', '3万元', '+50%', 'ROI下降约10%'],
        ['运营成本', '3万元', '-50%', 'ROI上升约10%'],
    ]
    add_table(doc, data[1:], data[0])
    
    doc.add_page_break()
    
    add_title(doc, '第六章 结论与建议', level=1)
    
    add_title(doc, '6.1 可行性结论', level=2)
    add_paragraph(doc, '经过全面的可行性分析，本项目在技术、经济、运营和法律四个维度均完全可行：', indent=True)
    
    conclusions = [
        '技术可行性：采用成熟的技术栈，所有核心技术均有明确解决方案，技术风险可控',
        '经济可行性：初期投入约4-7万元，年运营成本约2-5万元，投资回报率约70-200%',
        '运营可行性：运维难度低，用户接受度高，适合小型团队或个人运营',
        '法律可行性：数据采集合规，用户隐私保护到位，无重大法律风险',
        '综合评估：项目高度可行，建议立即启动第三阶段开发',
    ]
    add_list(doc, conclusions)
    
    add_title(doc, '6.2 开发建议', level=2)
    suggestions = [
        '按计划推进第三阶段开发，重点优化翻译质量和数据去重功能',
        '优先扩展数据源，增加更多权威科技媒体来源',
        '建立完善的日志系统和监控机制，便于问题排查',
        '定期收集用户反馈，持续优化产品体验',
        '关注技术趋势，适时引入新的技术方案（如大语言模型）',
        '考虑商业化方向，探索增值服务和盈利模式',
    ]
    add_list(doc, suggestions)
    
    add_title(doc, '6.3 风险提示', level=2)
    risks = [
        '爬虫行为需遵守目标网站的 robots.txt 协议，避免被封禁',
        '翻译服务免费额度有限，大规模使用可能需要付费',
        '新闻内容版权归原媒体所有，仅限个人/内部参考使用',
        '定期备份数据，防止数据丢失',
        '关注政策法规变化，确保合规运营',
    ]
    add_list(doc, risks)
    
    add_title(doc, '6.4 总结', level=2)
    add_paragraph(doc, '本项目是一个技术可行、经济合理、运营便捷、法律合规的 AI 新闻自动化处理平台。项目初期投入低、回报周期短、风险可控，具有良好的投资价值和发展前景。建议立即启动第三阶段开发工作，持续优化和扩展功能，提升用户体验和商业价值。', indent=True)
    
    doc.add_page_break()
    
    p = doc.add_paragraph()
    run = p.add_run('— 报告结束 —')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = os.path.join(os.path.dirname(__file__), 'AI新闻爬虫系统可行性分析与开发计划报告.docx')
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_report()
